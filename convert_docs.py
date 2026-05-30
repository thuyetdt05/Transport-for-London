# -*- coding: utf-8 -*-
"""
Script chuyển đổi file tong_hop_ky_thuat.md thành tong_hop_ky_thuat.docx và tong_hop_ky_thuat.pdf.
Sử dụng python-docx để tạo file DOCX chất lượng cao và xhtml2pdf để tạo file PDF chuyên nghiệp.
"""

import os
import re
import sys
from pathlib import Path

# Đảm bảo stdout/stderr encode UTF-8 để không bị lỗi tiếng Việt khi chạy trên console Windows
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')
if hasattr(sys.stderr, 'reconfigure'):
    sys.stderr.reconfigure(encoding='utf-8')

# Import các thư viện cần thiết
try:
    import docx
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("[Info] Cài đặt python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

try:
    import markdown
    from bs4 import BeautifulSoup
    from xhtml2pdf import pisa
except ImportError:
    print("[Info] Cài đặt markdown, beautifulsoup4, xhtml2pdf...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "markdown", "beautifulsoup4", "xhtml2pdf"])
    import markdown
    from bs4 import BeautifulSoup
    from xhtml2pdf import pisa


# ============================================================
# HELPER FUNCTIONS CHO DOCX
# ============================================================

def set_page_margins(doc, top=2, bottom=2, left=3, right=2):
    """Thiết lập lề trang theo cm (chuẩn báo cáo Việt Nam)"""
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def set_run_font(run, name="Arial", size=12, bold=False, italic=False, color=None):
    """Cấu hình font cho một run, hỗ trợ hiển thị tiếng Việt chính xác"""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    
    # Ép buộc font tiếng Việt trong Word XML
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    rPr.append(rFonts)


def add_hyperlink(paragraph, url, text, font_name="Arial", font_size=12, color="1F4E79", underline=True):
    """Thêm liên kết (hyperlink) vào đoạn văn DOCX"""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Font name
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)

    # Font size (nửa điểm)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size * 2)))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(font_size * 2)))
    rPr.append(szCs)

    # Color
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    # Underline
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink


def set_cell_background(cell, hex_color):
    """Thiết lập màu nền cho ô của bảng"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_all_borders(cell, hex_color="E2E8F0", size="4"):
    """Thiết lập toàn bộ viền cho ô"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), hex_color)
        tcBorders.append(b)
        
    tcPr.append(tcBorders)


def set_cell_left_border(cell, hex_color="1F4E79", size="24"):
    """Thiết lập viền trái dày và xóa các viền khác (dùng cho blockquote)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    # Viền trái
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), size)
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), hex_color)
    tcBorders.append(left)
    
    # Xóa viền còn lại
    for border_name in ['top', 'bottom', 'right']:
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'nil')
        tcBorders.append(b)
        
    tcPr.append(tcBorders)


# ============================================================
# PARSERS CHO MARKDOWN TRONG DOCX
# ============================================================

def parse_inline(text):
    """Parse text markdown chứa inline styles (bold, italic, code, links) thành các token"""
    pattern = re.compile(
        r'(?P<link>\[(?P<link_text>[^\]]+)\]\((?P<link_url>[^)]+)\))|'
        r'(?P<bold>\*\*(?P<bold_text>[^*]+)\*\*)|'
        r'(?P<code>`(?P<code_text>[^`]+)`)|'
        r'(?P<italic>\*(?P<italic_text>[^*]+)\*)|'
        r'(?P<italic2>_(?P<italic2_text>[^_]+)_)'
    )
    
    runs = []
    last_idx = 0
    for match in pattern.finditer(text):
        if match.start() > last_idx:
            runs.append(('text', text[last_idx:match.start()]))
        
        gd = match.groupdict()
        if gd['link']:
            runs.append(('link', gd['link_text'], gd['link_url']))
        elif gd['bold']:
            runs.append(('bold', gd['bold_text']))
        elif gd['code']:
            runs.append(('code', gd['code_text']))
        elif gd['italic']:
            runs.append(('italic', gd['italic_text']))
        elif gd['italic2']:
            runs.append(('italic', gd['italic2_text']))
            
        last_idx = match.end()
        
    if last_idx < len(text):
        runs.append(('text', text[last_idx:]))
        
    return runs


def add_styled_text(paragraph, text, default_font="Arial", default_size=12, default_color=None):
    """Thêm text có định dạng vào đoạn văn"""
    runs_data = parse_inline(text)
    for r_type, r_val, *extra in runs_data:
        if r_type == 'text':
            run = paragraph.add_run(r_val)
            set_run_font(run, name=default_font, size=default_size, color=default_color)
        elif r_type == 'bold':
            run = paragraph.add_run(r_val)
            set_run_font(run, name=default_font, size=default_size, bold=True, color=default_color)
        elif r_type == 'italic':
            run = paragraph.add_run(r_val)
            set_run_font(run, name=default_font, size=default_size, italic=True, color=default_color)
        elif r_type == 'code':
            run = paragraph.add_run(r_val)
            set_run_font(run, name="Consolas", size=default_size - 1, color=(199, 37, 78))
        elif r_type == 'link':
            url = extra[0]
            add_hyperlink(paragraph, url, r_val, font_name=default_font, font_size=default_size)


def parse_markdown(filepath):
    """Phân tách file markdown thành các block elements có cấu trúc"""
    with open(filepath, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    elements = []
    i = 0
    n = len(lines)
    
    while i < n:
        line = lines[i].rstrip('\r\n')
        
        # 1. Dòng trống
        if not line.strip():
            # Chỉ thêm nếu phần tử trước không phải dòng trống để tránh khoảng cách quá lớn
            if not elements or elements[-1][0] != 'blank':
                elements.append(('blank', ''))
            i += 1
            continue
            
        # 2. Heading
        if line.startswith('#'):
            m = re.match(r'^(#+)\s+(.*)$', line)
            if m:
                level = len(m.group(1))
                text = m.group(2)
                elements.append(('heading', level, text))
                i += 1
                continue
                
        # 3. Đường chia ngang (Horizontal rule)
        if line.strip() == '---':
            elements.append(('hr', ''))
            i += 1
            continue
            
        # 4. Blockquote
        if line.startswith('>'):
            bq_lines = []
            while i < n and lines[i].startswith('>'):
                bq_line = lines[i].rstrip('\r\n')
                # Bỏ ký tự '>' ở đầu dòng
                m = re.match(r'^>\s?(.*)$', bq_line)
                if m:
                    bq_lines.append(m.group(1))
                else:
                    bq_lines.append(bq_line[1:])
                i += 1
            elements.append(('blockquote', '\n'.join(bq_lines)))
            continue
            
        # 5. Code block
        if line.strip().startswith('```'):
            lang = line.strip()[3:].strip()
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip('\r\n'))
                i += 1
            if i < n:
                i += 1 # skip closing ```
            elements.append(('code_block', lang, '\n'.join(code_lines)))
            continue
            
        # 6. Table
        if line.strip().startswith('|'):
            table_lines = []
            while i < n and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].rstrip('\r\n'))
                i += 1
                
            if len(table_lines) >= 2:
                # Kiểm tra xem dòng 2 có phải là ngăn cách (separator) không
                sep_match = re.match(r'^\|\s*[:\-|\s]+\s*\|$', table_lines[1].strip())
                if sep_match:
                    headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
                    rows = []
                    for r_line in table_lines[2:]:
                        rows.append([c.strip() for c in r_line.split('|')[1:-1]])
                    elements.append(('table', headers, rows))
                else:
                    # Nếu không phải table hợp lệ, parse như paragraph
                    for tl in table_lines:
                        elements.append(('paragraph', tl))
            else:
                for tl in table_lines:
                    elements.append(('paragraph', tl))
            continue
            
        # 7. Unordered list item
        if line.strip().startswith('- ') or line.strip().startswith('* ') or line.strip().startswith('• '):
            indent = len(line) - len(line.lstrip())
            # Cắt bỏ ký tự bullet
            text = line.strip()[2:]
            elements.append(('ul_item', indent, text))
            i += 1
            continue
            
        # 8. Ordered list item
        m_ol = re.match(r'^(\s*)(\d+)\.\s+(.*)$', line)
        if m_ol:
            indent = len(m_ol.group(1))
            num = m_ol.group(2)
            text = m_ol.group(3)
            elements.append(('ol_item', indent, num, text))
            i += 1
            continue
            
        # 9. Paragraph thông thường
        para_lines = [line]
        i += 1
        # Gom các dòng chữ liên tiếp lại với nhau cho đến khi gặp dòng trống hoặc block khác
        while i < n:
            next_line = lines[i].rstrip('\r\n')
            if not next_line.strip():
                break
            # Dấu hiệu bắt đầu block khác
            if (next_line.startswith('#') or 
                next_line.strip() == '---' or 
                next_line.startswith('>') or 
                next_line.strip().startswith('```') or 
                next_line.strip().startswith('|') or 
                next_line.strip().startswith('- ') or 
                next_line.strip().startswith('* ') or 
                re.match(r'^(\s*)(\d+)\.\s+(.*)$', next_line)):
                break
            para_lines.append(next_line)
            i += 1
        elements.append(('paragraph', ' '.join(para_lines)))
        
    return elements


# ============================================================
# CHUYỂN ĐỔI SANG DOCX
# ============================================================

def convert_md_to_docx(md_filepath, output_docx_path):
    print(f"[Info] Đang chuyển đổi Markdown sang DOCX: {output_docx_path}...")
    elements = parse_markdown(md_filepath)
    
    doc = Document()
    set_page_margins(doc, top=2, bottom=2, left=3, right=2)
    
    # Thiết lập khoảng cách dòng mặc định
    doc.styles['Normal'].font.name = 'Arial'
    
    for item in elements:
        el_type = item[0]
        
        if el_type == 'blank':
            # Dòng trống
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            
        elif el_type == 'heading':
            level, text = item[1], item[2]
            
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            
            # Map heading levels (MD # -> Doc Title, ## -> H1, ### -> H2, #### -> H3)
            if level == 1:
                # Tiêu đề chính
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(12)
                # Loại bỏ emoji nếu có để Word trông chuẩn mực hơn, hoặc giữ nguyên
                # Ở đây chúng ta giữ nguyên text
                run = p.add_run(text)
                set_run_font(run, name="Arial", size=16, bold=True, color=(31, 78, 121)) # Hex #1F4E79
            elif level == 2:
                # H1
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(text)
                set_run_font(run, name="Arial", size=13, bold=True, color=(31, 78, 121))
                # Vẽ một đường gạch dưới nhỏ sau H1 để trang trí nếu muốn (hoặc bỏ qua)
            elif level == 3:
                # H2
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(text)
                set_run_font(run, name="Arial", size=12, bold=True, italic=True, color=(51, 65, 85)) # #334155
            else:
                # H3 trở xuống
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(text)
                set_run_font(run, name="Arial", size=11, bold=True, color=(51, 65, 85))
                
        elif el_type == 'hr':
            # Vẽ đường kẻ ngang trang trí
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4') # Kích thước nét kẻ
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'E2E8F0') # Xám nhạt
            pBdr.append(bottom)
            pPr.append(pBdr)
            
        elif el_type == 'blockquote':
            text = item[1]
            # Blockquote được mô phỏng bằng một bảng 1 ô có viền trái dày màu xanh
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            
            cell = table.rows[0].cells[0]
            # Thiết lập chiều rộng khoảng 15.5cm (bằng độ rộng của vùng in ấn 21 - 3 - 2 = 16)
            cell.width = Cm(15.5)
            
            set_cell_background(cell, 'F8FAFC') # Nền xám nhạt
            set_cell_left_border(cell, '1F4E79', '24') # Viền trái 3pt xanh đậm
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.2
            
            # Text bên trong blockquote viết nghiêng và nhỏ hơn chút
            add_styled_text(p, text.strip(), default_font="Arial", default_size=11, default_color=(71, 85, 105)) # Nghiêng
            # Làm xiên các chữ bằng tay (vì add_styled_text sẽ đọc các token, nhưng chúng ta muốn cả đoạn nghiêng)
            for r in p.runs:
                r.italic = True
                
        elif el_type == 'code_block':
            lang, code = item[1], item[2]
            # Code block được đặt trong 1 bảng 1 ô viền mỏng màu xám nhạt
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            
            cell = table.rows[0].cells[0]
            cell.width = Cm(15.5)
            
            set_cell_background(cell, 'F8FAFC')
            set_cell_all_borders(cell, 'E2E8F0', '4') # Viền xám mỏng
            
            # Ghi mã code dòng-dòng
            # Vì mặc định đã có một paragraph sẵn trong cell, ta dùng nó trước
            p_first = cell.paragraphs[0]
            p_first.paragraph_format.space_before = Pt(3)
            p_first.paragraph_format.space_after = Pt(1)
            p_first.paragraph_format.line_spacing = 1.0 # Dòng đơn cho code
            
            code_lines = code.split('\n')
            p_first.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p_first.add_run(code_lines[0])
            set_run_font(run, name="Consolas", size=9.5, color=(30, 41, 59))
            
            for c_line in code_lines[1:]:
                p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.0
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(c_line)
                set_run_font(run, name="Consolas", size=9.5, color=(30, 41, 59))
                
            # Thêm lề dưới cho dòng cuối cùng
            cell.paragraphs[-1].paragraph_format.space_after = Pt(3)
            
        elif el_type == 'table':
            headers, rows = item[1], item[2]
            create_docx_table(doc, headers, rows)
            
        elif el_type == 'ul_item':
            indent, text = item[1], item[2]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            # Thụt lề phụ thuộc vào độ thụt dòng trong markdown (indent // 2 hoặc tương tự)
            indent_level = indent // 2
            p.paragraph_format.left_indent = Cm(0.5 + 0.4 * indent_level)
            
            add_styled_text(p, text, default_font="Arial", default_size=12)
            
        elif el_type == 'ol_item':
            indent, num, text = item[1], item[2], item[3]
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            indent_level = indent // 2
            p.paragraph_format.left_indent = Cm(0.5 + 0.4 * indent_level)
            
            # Thêm nội dung
            add_styled_text(p, text, default_font="Arial", default_size=12)
            
        elif el_type == 'paragraph':
            text = item[1]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.3
            
            add_styled_text(p, text, default_font="Arial", default_size=12)
            
    doc.save(output_docx_path)
    print("[Success] Đã tạo thành công file DOCX!")


def create_docx_table(doc, headers, rows):
    """Helper tạo bảng DOCX từ dữ liệu parsed"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    
    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_background(cell, '1F4E79') # Nền xanh lục đậm
        set_cell_all_borders(cell, 'CBD5E1', '4')
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        
        run = p.add_run(h)
        set_run_font(run, name="Arial", size=10, bold=True, color=(255, 255, 255))
        
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.add_row()
        bg_color = 'F8FAFC' if r_idx % 2 == 0 else 'FFFFFF' # Zebra striping
        for c_idx, val in enumerate(row_data):
            if c_idx >= len(row.cells):
                break
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_background(cell, bg_color)
            set_cell_all_borders(cell, 'E2E8F0', '4')
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            
            add_styled_text(p, val.strip(), default_font="Arial", default_size=9.5)
            
    # Thêm một paragraph trống nhỏ để tạo khoảng trống sau bảng
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(6)
    p_after.paragraph_format.space_after = Pt(6)


# ============================================================
# CHUYỂN ĐỔI SANG PDF
# ============================================================

def convert_md_to_pdf(md_filepath, output_pdf_path):
    print(f"[Info] Đang chuyển đổi Markdown sang PDF: {output_pdf_path}...")
    
    # 1. Đọc nội dung Markdown
    with open(md_filepath, "r", encoding="utf-8") as f:
        md_content = f.read()
        
    # 2. Convert Markdown sang HTML thông qua thư viện markdown
    # Sử dụng các extensions: extra (để parse tables), fenced_code (parse code blocks)
    html_content = markdown.markdown(md_content, extensions=['extra', 'fenced_code'])
    
    # 3. Post-process HTML để chèn các class cho table zebra striping và chuẩn hóa
    soup = BeautifulSoup(html_content, 'html.parser')
    for table in soup.find_all('table'):
        rows = table.find_all('tr')
        for idx, row in enumerate(rows):
            if idx == 0:
                continue # Bỏ qua header
            if idx % 2 == 0:
                row['class'] = row.get('class', []) + ['even']
                
    processed_html = str(soup)
    
    # 4. Đăng ký font Arial hệ thống trực tiếp với ReportLab để hỗ trợ tiếng Việt
    # Việc đăng ký trực tiếp ở mức ReportLab và sử dụng registerFontFamily giúp tránh
    # lỗi PermissionError / URL parser khi xhtml2pdf xử lý các thẻ @font-face trong CSS.
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfbase.pdfmetrics import registerFontFamily
    
    sys_root = os.environ.get("SystemRoot", "C:\\Windows")
    arial_path = os.path.join(sys_root, "Fonts", "arial.ttf")
    arial_bold_path = os.path.join(sys_root, "Fonts", "arialbd.ttf")
    arial_italic_path = os.path.join(sys_root, "Fonts", "ariali.ttf")
    arial_bi_path = os.path.join(sys_root, "Fonts", "arialbi.ttf")
    
    if os.path.exists(arial_path):
        try:
            pdfmetrics.registerFont(TTFont("Arial", arial_path))
            pdfmetrics.registerFont(TTFont("Arial-Bold", arial_bold_path))
            pdfmetrics.registerFont(TTFont("Arial-Italic", arial_italic_path))
            pdfmetrics.registerFont(TTFont("Arial-BoldItalic", arial_bi_path))
            registerFontFamily("Arial", normal="Arial", bold="Arial-Bold", italic="Arial-Italic", boldItalic="Arial-BoldItalic")
            font_family_css = "Arial, sans-serif"
            print("[Info] Đã đăng ký thành công font Arial hệ thống.")
        except Exception as e:
            print(f"[Warning] Không thể đăng ký font Arial: {e}. Sẽ sử dụng font mặc định.")
            font_family_css = "Helvetica, sans-serif"
    else:
        font_family_css = "Helvetica, sans-serif"
        
    # 5. Tạo mẫu HTML hoàn chỉnh với CSS stylesheet chất lượng cao
    # Style được thiết kế đặc thù cho xhtml2pdf nhằm hiển thị tiếng Việt bằng font Arial hệ thống
    # Không dùng @font-face trong CSS nữa, dùng trực tiếp font-family Arial đã được đăng ký
    full_html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
@page {{
    size: a4;
    margin: 2cm 2cm 2cm 3cm; /* lề trên, phải, dưới, trái */
}}
body {{
    font-family: {font_family_css};
    font-size: 11pt;
    line-height: 1.5;
    color: #1E293B;
}}
h1 {{
    font-size: 18pt;
    color: #1F4E79;
    text-align: center;
    margin-top: 10px;
    margin-bottom: 20px;
}}
h2 {{
    font-size: 13pt;
    color: #1F4E79;
    border-bottom: 0.5px solid #CBD5E1;
    padding-bottom: 3px;
    margin-top: 20px;
    margin-bottom: 10px;
}}
h3 {{
    font-size: 11.5pt;
    color: #334155;
    margin-top: 15px;
    margin-bottom: 6px;
}}
p {{
    margin-top: 0;
    margin-bottom: 8px;
    text-align: justify;
}}
blockquote {{
    margin: 10px 0;
    padding: 8px 12px;
    background-color: #F8FAFC;
    border-left: 4px solid #1F4E79;
    color: #475569;
    font-style: italic;
}}
pre {{
    background-color: #F8FAFC;
    border: 0.5px solid #E2E8F0;
    padding: 8px;
    margin-top: 5px;
    margin-bottom: 10px;
    font-family: "Courier New", monospace;
    font-size: 9pt;
}}
code {{
    font-family: "Courier New", monospace;
    font-size: 9pt;
    color: #C7254E;
    background-color: #F9F2F4;
    padding: 1px 3px;
}}
ul, ol {{
    margin-bottom: 8px;
    padding-left: 20px;
}}
li {{
    margin-bottom: 3px;
}}
table {{
    width: 100%;
    border-collapse: collapse;
    margin-top: 12px;
    margin-bottom: 15px;
    font-size: 9.5pt;
}}
th {{
    background-color: #1F4E79;
    color: white;
    font-weight: bold;
    border: 0.5px solid #CBD5E1;
    padding: 6px;
    text-align: center;
}}
td {{
    border: 0.5px solid #E2E8F0;
    padding: 6px;
    text-align: left;
}}
tr.even td {{
    background-color: #F8FAFC;
}}
hr {{
    border: 0;
    border-top: 0.5px solid #CBD5E1;
    margin: 18px 0;
}}
</style>
</head>
<body>
{processed_html}
</body>
</html>"""

    # 6. Thực hiện convert thông qua xhtml2pdf pisa
    with open(output_pdf_path, "w+b") as result_file:
        pisa_status = pisa.CreatePDF(
            src=full_html,
            dest=result_file,
            encoding='utf-8'
        )
        
    if not pisa_status.err:
        print("[Success] Đã tạo thành công file PDF!")
    else:
        print(f"[Error] Có lỗi xảy ra khi tạo PDF: {pisa_status.err}")


# ============================================================
# MAIN EXECUTION
# ============================================================

if __name__ == "__main__":
    # Đường dẫn nguồn và đích
    workspace_dir = Path(r"c:\Users\My PC\Downloads\DE")
    md_file = workspace_dir / "tong_hop_ky_thuat.md"
    docx_file = workspace_dir / "tong_hop_ky_thuat.docx"
    pdf_file = workspace_dir / "tong_hop_ky_thuat.pdf"
    
    if not md_file.exists():
        print(f"[Error] Không tìm thấy file markdown tại: {md_file}")
        sys.exit(1)
        
    print(f"Bắt đầu chuyển đổi tệp {md_file}...")
    
    try:
        convert_md_to_docx(md_file, docx_file)
    except Exception as e:
        print(f"[Error] Lỗi khi tạo DOCX: {e}")
        import traceback
        traceback.print_exc()
        
    try:
        convert_md_to_pdf(md_file, pdf_file)
    except Exception as e:
        print(f"[Error] Lỗi khi tạo PDF: {e}")
        import traceback
        traceback.print_exc()
        
    print("Hoàn thành quy trình chuyển đổi!")
