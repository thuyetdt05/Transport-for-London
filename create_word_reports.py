# -*- coding: utf-8 -*-
"""
Script tạo 2 file Word báo cáo cho dự án TfL London - Nhóm 14
"""

import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ============================================================
# HELPERS
# ============================================================

def set_page_margins(doc, top=2, bottom=2, left=3, right=2):
    """Set page margins in cm"""
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def set_run_font(run, name="Times New Roman", size=13, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Force Vietnamese font
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    rPr.append(rFonts)


def add_paragraph(doc, text="", style=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  size=13, bold=False, italic=False, space_before=0, space_after=6,
                  line_spacing=1.5, color=None, first_line_indent=None):
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    
    pf = p.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    
    return p


def add_heading_chapter(doc, number, title, size=14):
    """Add chapter heading - bold, centered"""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f"CHƯƠNG {number}: {title}")
    set_run_font(run, size=size, bold=True)
    return p


def add_heading_1(doc, numbering, title, size=13):
    """Add level 1 heading like 1.1, 2.3..."""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f"{numbering} {title}")
    set_run_font(run, size=size, bold=True)
    return p


def add_heading_2(doc, numbering, title, size=13):
    """Add level 2 heading like 1.1.1"""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f"{numbering} {title}")
    set_run_font(run, size=size, bold=True, italic=True)
    return p


def add_heading_3(doc, numbering, title, size=13):
    """Add level 3 heading like 1.1.1.1"""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f"{numbering} {title}")
    set_run_font(run, size=13, bold=True, italic=False)
    return p


def add_body_text(doc, text, indent=True):
    """Add justified body text with first line indent"""
    p = add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                      size=13, space_after=6, line_spacing=1.5,
                      first_line_indent=1.0 if indent else None)
    return p


def add_table_with_caption(doc, caption_text, headers, data_rows, col_widths=None):
    """Add a table with caption above"""
    # Caption above table
    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(3)
    run = p_cap.add_run(caption_text)
    set_run_font(run, size=12, bold=True, italic=True)
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=12, bold=True)
        # Header background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F4E79')
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for row_data in data_rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_run_font(run, size=12)
    
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    
    # Space after table
    doc.add_paragraph()
    return table


def add_figure_caption(doc, caption_text):
    """Add figure caption below image"""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(9)
    run = p.add_run(caption_text)
    set_run_font(run, size=12, bold=False, italic=True)
    return p


def add_page_break(doc):
    doc.add_page_break()


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F4E79')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ============================================================
# FILE 1: BÁO CÁO CHÍNH THỨC
# ============================================================

def create_main_report():
    doc = Document()
    set_page_margins(doc, top=2, bottom=2, left=3, right=2)
    
    # ====================
    # TRANG BÌA
    # ====================
    
    # Tiêu đề trường
    for text, sz, bold in [
        ("BỘ CÔNG THƯƠNG", 13, True),
        ("TRƯỜNG ĐẠI HỌC CÔNG NGHIỆP THÀNH PHỐ HỒ CHÍ MINH", 14, True),
        ("KHOA CÔNG NGHỆ THÔNG TIN", 13, True),
    ]:
        p = add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, size=sz, bold=bold, space_after=4)
    
    add_horizontal_line(doc)
    
    # Khoảng cách
    for _ in range(3):
        add_paragraph(doc, "")
    
    # Loại báo cáo
    p = add_paragraph(doc, "BÁO CÁO BÀI TẬP LỚN", align=WD_ALIGN_PARAGRAPH.CENTER,
                      size=14, bold=True, space_after=4)
    p = add_paragraph(doc, "Môn học: Nhập môn Kỹ thuật Dữ liệu", align=WD_ALIGN_PARAGRAPH.CENTER,
                      size=13, bold=False, space_after=4)
    
    for _ in range(2):
        add_paragraph(doc, "")
    
    # Tên đề tài
    p = add_paragraph(doc, "ĐỀ TÀI:", align=WD_ALIGN_PARAGRAPH.CENTER,
                      size=15, bold=True, space_after=6)
    p = add_paragraph(doc, "LONDON TRANSPORT ANALYSIS", align=WD_ALIGN_PARAGRAPH.CENTER,
                      size=18, bold=True, color=(31, 78, 121), space_after=4)
    p = add_paragraph(doc, "Phân tích Hệ thống Giao thông Công cộng London (TfL)", align=WD_ALIGN_PARAGRAPH.CENTER,
                      size=14, bold=False, italic=True, space_after=8)
    
    for _ in range(2):
        add_paragraph(doc, "")
    
    # Thông tin nhóm - dùng bảng
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    info_data = [
        ("Nhóm thực hiện:", "Nhóm 14"),
        ("Thành viên:", "Nguyễn Văn A - MSSV: 23XXXX01\nNguyễn Văn B - MSSV: 23XXXX02\nNguyễn Văn C - MSSV: 23XXXX03\nNguyễn Văn D - MSSV: 23XXXX04"),
        ("Lớp:", "DHKHDL20A"),
        ("Giảng viên hướng dẫn:", "TS. Lê Trọng Ngọc"),
        ("Học kỳ:", "II - Năm học 2025–2026"),
        ("Ngày nộp:", "Tháng 05/2026"),
    ]
    
    for i, (label, value) in enumerate(info_data):
        row = table.rows[i]
        # Label cell
        c0 = row.cells[0]
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r0 = p0.add_run(label)
        set_run_font(r0, size=13, bold=True)
        c0.width = Cm(6)
        
        # Value cell
        c1 = row.cells[1]
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(value)
        set_run_font(r1, size=13)
        c1.width = Cm(9)
    
    for _ in range(4):
        add_paragraph(doc, "")
    
    add_paragraph(doc, "TP. HỒ CHÍ MINH, THÁNG 05 NĂM 2026",
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True)
    
    add_page_break(doc)
    
    # ====================
    # LỜI CẢM ƠN
    # ====================
    add_paragraph(doc, "LỜI CẢM ƠN", align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=14, bold=True, space_after=12)
    
    cam_on_text = [
        "Để hoàn thành báo cáo bài tập lớn này, nhóm chúng em xin gửi lời cảm ơn chân thành và sâu sắc đến Tiến sĩ Lê Trọng Ngọc — giảng viên phụ trách môn học Nhập môn Kỹ thuật Dữ liệu, Khoa Công nghệ Thông tin, Trường Đại học Công nghiệp Thành phố Hồ Chí Minh. Với sự hướng dẫn tận tình, những góp ý chuyên môn sâu sắc và tinh thần khuyến khích học tập không ngừng của Thầy, nhóm chúng em đã có nền tảng vững chắc để triển khai một dự án kỹ thuật dữ liệu mang tính thực tiễn cao.",
        "Chúng em cũng gửi lời tri ân đến Ban Giám hiệu nhà trường, Ban Chủ nhiệm Khoa Công nghệ Thông tin và toàn thể quý thầy cô đã tạo điều kiện về cơ sở vật chất, tài nguyên học tập và môi trường nghiên cứu thuận lợi trong suốt quá trình học tập tại trường.",
        "Dự án này sử dụng nguồn dữ liệu mở từ Transport for London (TfL) và National Public Transport Access Nodes (NaPTAN) — hai tổ chức đã và đang đóng góp to lớn cho sự minh bạch thông tin trong lĩnh vực giao thông công cộng. Nhóm trân trọng sự công khai hóa dữ liệu này, tạo cơ hội để sinh viên tiếp cận và thực hành trên dữ liệu thực tế quy mô lớn.",
        "Bên cạnh đó, nhóm xin cảm ơn toàn thể cộng đồng mã nguồn mở đã phát triển và duy trì các thư viện Python như Pandas, NumPy, Scikit-learn, Folium, SQLite3 — những công cụ đã hỗ trợ đắc lực trong quá trình xây dựng pipeline ETL và hệ thống phân tích dữ liệu.",
        "Cuối cùng, chúng em xin cảm ơn gia đình và bạn bè đã luôn ủng hộ, động viên tinh thần trong suốt quá trình thực hiện dự án. Mặc dù đã cố gắng hết sức, báo cáo chắc chắn còn những thiếu sót nhất định. Nhóm rất mong nhận được sự góp ý từ Thầy và Hội đồng để hoàn thiện hơn trong các dự án tiếp theo.",
        "Nhóm 14 xin chân thành cảm ơn!",
    ]
    
    for text in cam_on_text:
        add_body_text(doc, text)
    
    add_paragraph(doc, "")
    p = add_paragraph(doc, "                                                        TP. HCM, tháng 05 năm 2026",
                      align=WD_ALIGN_PARAGRAPH.LEFT, size=13, italic=True)
    add_paragraph(doc, "                                                        Nhóm 14 — Lớp DHKHDL20A",
                  align=WD_ALIGN_PARAGRAPH.LEFT, size=13, bold=True)
    
    add_page_break(doc)
    
    # ====================
    # TÓM TẮT (ABSTRACT)
    # ====================
    add_paragraph(doc, "TÓM TẮT", align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=14, bold=True, space_after=6)
    add_paragraph(doc, "ABSTRACT", align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=13, bold=True, italic=True, space_after=12)
    
    abstract_vn = [
        "Hệ thống giao thông công cộng của thành phố London do Transport for London (TfL) quản lý và vận hành là một trong những mạng lưới giao thông đô thị phức tạp và có quy mô lớn nhất thế giới, phục vụ hàng triệu lượt hành khách mỗi ngày thông qua nhiều phương thức vận chuyển khác nhau, bao gồm tàu điện ngầm (Underground), tàu trên cao (Overground), tàu điện nhẹ DLR, tuyến Elizabeth và hệ thống xe điện (Tram).",
        "Báo cáo này trình bày toàn bộ quá trình xây dựng một pipeline kỹ thuật dữ liệu (ETL — Extract, Transform, Load) hoàn chỉnh nhằm phân tích hệ thống nhà ga TfL trong giai đoạn 2017–2021. Dự án được thực hiện bởi Nhóm 14, sinh viên môn học Nhập môn Kỹ thuật Dữ liệu, Khoa Công nghệ Thông tin, Trường Đại học Công nghiệp TP.HCM.",
        "Quá trình Extract thu thập và tích hợp dữ liệu từ ba nguồn độc lập: (1) file KML chứa tọa độ địa lý của các nhà ga, (2) file CSV từ TfL chứa lưu lượng hành khách Enter/Exit theo từng năm, và (3) file Stops.csv từ cơ sở dữ liệu NaPTAN chứa metadata về điểm dừng giao thông. Giai đoạn Transform áp dụng các kỹ thuật làm sạch dữ liệu nâng cao như chuẩn hóa tên ga bằng Regex, thuật toán Nearest Neighbor để ghép tọa độ, Spatial Join để enrich thông tin địa lý, và Feature Engineering để tạo ra các đặc trưng phân tích.",
        "Thuật toán phân cụm KMeans với K=6 được áp dụng trên tập đặc trưng gồm lưu lượng hành khách năm 2021, số tuyến phục vụ và tọa độ địa lý, phân loại 298 nhà ga thành 6 cụm có đặc trưng riêng biệt: Siêu trung tâm (8 ga, trung bình 37,3 triệu lượt/năm), Ga lớn (42 ga), Ga trung bình (78 ga), Ga nhỏ (80 ga), Ga ít khách (37 ga) và Ga rất ít khách (53 ga).",
        "Phân tích tác động COVID-19 cho thấy tổng lưu lượng hành khách giảm mạnh 63,21% từ năm 2019 (3,28 tỷ lượt) xuống còn 1,21 tỷ lượt vào năm 2020. Năm 2021 ghi nhận dấu hiệu phục hồi tích cực với mức tăng 16,82% so với năm 2020. Kết quả trực quan hóa được thể hiện thông qua bản đồ tương tác Folium với đầy đủ tính năng lọc theo cụm, tìm kiếm ga, bản đồ nhiệt và biểu đồ thống kê động.",
        "Dữ liệu được lưu trữ đa định dạng vào SQLite (cơ sở dữ liệu quan hệ với bảng fact_stations và dim_clusters), Excel và CSV để phục vụ các mục đích phân tích khác nhau. Hệ thống web server tích hợp (serve_outputs.py) cho phép chia sẻ kết quả qua mạng cục bộ và internet thông qua các công nghệ tunnel như ngrok và localtunnel.",
        "Kết quả của dự án không chỉ cung cấp góc nhìn toàn diện về mạng lưới giao thông công cộng London mà còn xây dựng được một framework ETL có khả năng tái sử dụng và mở rộng cho các bài toán phân tích dữ liệu giao thông đô thị tương tự.",
    ]
    
    for text in abstract_vn:
        add_body_text(doc, text)
    
    add_paragraph(doc, "")
    add_paragraph(doc, "Từ khóa: Transport for London (TfL), ETL Pipeline, KMeans Clustering, COVID-19 Impact Analysis, Spatial Join, Folium Visualization, SQLite, NaPTAN.",
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, italic=True, space_after=6)
    
    add_page_break(doc)
    
    # ====================
    # DANH MỤC TỪ VIẾT TẮT
    # ====================
    add_paragraph(doc, "DANH MỤC TỪ VIẾT TẮT", align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=14, bold=True, space_after=12)
    
    abbreviations = [
        ("TfL", "Transport for London — Cơ quan quản lý giao thông London"),
        ("ETL", "Extract, Transform, Load — Quy trình trích xuất, biến đổi và tải dữ liệu"),
        ("KML", "Keyhole Markup Language — Định dạng XML mô tả dữ liệu địa lý"),
        ("CSV", "Comma-Separated Values — Định dạng dữ liệu phân tách bằng dấu phẩy"),
        ("NaPTAN", "National Public Transport Access Nodes — Cơ sở dữ liệu điểm dừng giao thông quốc gia Anh"),
        ("KMeans", "K-Means Clustering — Thuật toán phân cụm K trung tâm"),
        ("DLR", "Docklands Light Railway — Tuyến tàu điện nhẹ khu Docklands"),
        ("LU", "London Underground — Tàu điện ngầm London"),
        ("LO", "London Overground — Tàu trên cao London"),
        ("EL", "Elizabeth Line — Tuyến Elizabeth (Crossrail)"),
        ("GPS", "Global Positioning System — Hệ thống định vị toàn cầu"),
        ("API", "Application Programming Interface — Giao diện lập trình ứng dụng"),
        ("HTML", "HyperText Markup Language — Ngôn ngữ đánh dấu siêu văn bản"),
        ("JSON", "JavaScript Object Notation — Định dạng dữ liệu nhẹ"),
        ("DB", "Database — Cơ sở dữ liệu"),
        ("SQL", "Structured Query Language — Ngôn ngữ truy vấn có cấu trúc"),
        ("GVHD", "Giảng viên hướng dẫn"),
        ("TP.HCM", "Thành phố Hồ Chí Minh"),
        ("MSSV", "Mã số sinh viên"),
        ("COVID-19", "Coronavirus Disease 2019 — Dịch bệnh viêm đường hô hấp cấp do SARS-CoV-2"),
    ]
    
    add_table_with_caption(doc, "Bảng 0.1: Danh mục từ viết tắt được sử dụng trong báo cáo",
                           ["Từ viết tắt", "Ý nghĩa đầy đủ"], abbreviations,
                           col_widths=[3.5, 12])
    
    add_page_break(doc)
    
    # ====================
    # DANH MỤC BẢNG BIỂU
    # ====================
    add_paragraph(doc, "DANH MỤC BẢNG BIỂU", align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=14, bold=True, space_after=12)
    
    bang_bieu = [
        ("Bảng 0.1", "Danh mục từ viết tắt được sử dụng trong báo cáo"),
        ("Bảng 2.1", "Tổng quan ba nguồn dữ liệu đầu vào của dự án"),
        ("Bảng 2.2", "Cấu trúc dữ liệu file TfL_stations.csv"),
        ("Bảng 2.3", "Các loại StopType trong NaPTAN và mức độ ưu tiên"),
        ("Bảng 3.1", "Các thư viện và công cụ sử dụng trong dự án"),
        ("Bảng 3.2", "Cấu trúc bảng fact_stations trong SQLite"),
        ("Bảng 3.3", "Cấu trúc bảng dim_clusters trong SQLite"),
        ("Bảng 4.1", "Kết quả Regex chuẩn hóa tên ga — ví dụ minh họa"),
        ("Bảng 4.2", "Thống kê kết quả quá trình ghép dữ liệu (Merge)"),
        ("Bảng 4.3", "Các đặc trưng đầu vào cho mô hình KMeans"),
        ("Bảng 5.1", "Tổng hợp kết quả phân cụm KMeans 6 cụm"),
        ("Bảng 5.2", "Top 10 nhà ga đông khách nhất năm 2021"),
        ("Bảng 5.3", "Tác động COVID-19 theo từng cụm nhà ga"),
        ("Bảng 6.1", "So sánh ưu nhược điểm của hệ thống"),
    ]
    
    for so, ten in bang_bieu:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        r1 = p.add_run(f"{so}: ")
        set_run_font(r1, size=13, bold=True)
        r2 = p.add_run(ten)
        set_run_font(r2, size=13)
    
    add_page_break(doc)
    
    # ====================
    # DANH MỤC HÌNH ẢNH
    # ====================
    add_paragraph(doc, "DANH MỤC HÌNH ẢNH VÀ SƠ ĐỒ", align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=14, bold=True, space_after=12)
    
    hinh_anh = [
        ("Hình 3.1", "Kiến trúc tổng thể Pipeline ETL End-to-End"),
        ("Hình 3.2", "Sơ đồ quan hệ cơ sở dữ liệu SQLite"),
        ("Hình 4.1", "Quy trình chuẩn hóa tên ga bằng Regex"),
        ("Hình 4.2", "Minh họa thuật toán Nearest Neighbor Matching"),
        ("Hình 4.3", "Biểu đồ Elbow Method chọn K=6 cho KMeans"),
        ("Hình 4.4", "Bản đồ tương tác Folium — Tổng quan nhà ga TfL"),
        ("Hình 5.1", "Biểu đồ phân bố 6 cụm nhà ga trên bản đồ London"),
        ("Hình 5.2", "Biểu đồ so sánh lưu lượng hành khách 2017–2021"),
        ("Hình 5.3", "Biểu đồ tác động COVID-19 theo từng cụm"),
        ("Hình 5.4", "Bản đồ nhiệt (Heatmap) mật độ hành khách"),
        ("Hình 5.5", "Biểu đồ xu hướng 5 năm của các nhà ga"),
    ]
    
    for so, ten in hinh_anh:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        r1 = p.add_run(f"{so}: ")
        set_run_font(r1, size=13, bold=True)
        r2 = p.add_run(ten)
        set_run_font(r2, size=13)
    
    add_page_break(doc)
    
    # ====================
    # MỤC LỤC (thủ công)
    # ====================
    add_paragraph(doc, "MỤC LỤC", align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=14, bold=True, space_after=12)
    
    toc_items = [
        ("LỜI CẢM ƠN", "", True, False),
        ("TÓM TẮT", "", True, False),
        ("DANH MỤC TỪ VIẾT TẮT", "", True, False),
        ("DANH MỤC BẢNG BIỂU", "", True, False),
        ("DANH MỤC HÌNH ẢNH VÀ SƠ ĐỒ", "", True, False),
        ("CHƯƠNG 1: MỞ ĐẦU", "1", True, False),
        ("1.1 Lý do chọn đề tài", "1", False, False),
        ("1.2 Mục tiêu nghiên cứu", "2", False, False),
        ("1.3 Phạm vi và đối tượng nghiên cứu", "2", False, False),
        ("1.4 Ý nghĩa khoa học và thực tiễn", "3", False, False),
        ("1.5 Cấu trúc báo cáo", "4", False, False),
        ("CHƯƠNG 2: TỔNG QUAN VỀ DỮ LIỆU", "5", True, False),
        ("2.1 Giới thiệu về Transport for London (TfL)", "5", False, False),
        ("2.2 Mô tả chi tiết các nguồn dữ liệu", "6", False, False),
        ("2.3 Thách thức của dữ liệu thô", "9", False, False),
        ("2.4 Đánh giá chất lượng bộ dữ liệu", "10", False, False),
        ("CHƯƠNG 3: THIẾT KẾ HỆ THỐNG", "12", True, False),
        ("3.1 Kiến trúc Pipeline ETL End-to-End", "12", False, False),
        ("3.2 Thiết kế cơ sở dữ liệu SQLite", "14", False, False),
        ("3.3 Công nghệ, thư viện và công cụ sử dụng", "15", False, False),
        ("CHƯƠNG 4: TRIỂN KHAI VÀ THỰC HIỆN", "17", True, False),
        ("4.1 Extract — Thu thập dữ liệu từ nhiều nguồn", "17", False, False),
        ("4.2 Transform — Làm sạch và biến đổi dữ liệu", "19", False, False),
        ("4.3 Phân cụm KMeans", "22", False, False),
        ("4.4 Trực quan hóa tương tác với Folium", "24", False, False),
        ("4.5 Lưu trữ và xuất dữ liệu", "25", False, False),
        ("CHƯƠNG 5: KẾT QUẢ VÀ PHÂN TÍCH", "27", True, False),
        ("5.1 Kết quả phân cụm KMeans (6 cụm)", "27", False, False),
        ("5.2 Phân tích tác động COVID-19", "29", False, False),
        ("5.3 Các insight quan trọng và Top ga đông khách", "31", False, False),
        ("5.4 Đánh giá hiệu suất mô hình", "33", False, False),
        ("CHƯƠNG 6: ĐÁNH GIÁ HỆ THỐNG", "35", True, False),
        ("6.1 Ưu điểm của hệ thống", "35", False, False),
        ("6.2 Hạn chế và khó khăn gặp phải", "36", False, False),
        ("6.3 Giải pháp khắc phục", "37", False, False),
        ("CHƯƠNG 7: KẾT LUẬN VÀ KIẾN NGHỊ", "38", True, False),
        ("7.1 Kết luận", "38", False, False),
        ("7.2 Hướng phát triển tương lai", "39", False, False),
        ("TÀI LIỆU THAM KHẢO", "41", True, False),
        ("PHỤ LỤC", "43", True, False),
    ]
    
    for item_text, page_num, is_bold, _ in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.3
        indent = Cm(0) if is_bold else Cm(0.8)
        p.paragraph_format.left_indent = indent
        run = p.add_run(item_text)
        set_run_font(run, size=12, bold=is_bold)
        if page_num:
            tab_run = p.add_run(f"\t{page_num}")
            set_run_font(tab_run, size=12)
        # Tab stop at right
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:leader'), 'dot')
        tab.set(qn('w:pos'), '8640')
        tabs.append(tab)
        pPr.append(tabs)
    
    add_page_break(doc)
    
    # ====================
    # CHƯƠNG 1: MỞ ĐẦU
    # ====================
    add_heading_chapter(doc, "1", "MỞ ĐẦU")
    
    add_heading_1(doc, "1.1", "Lý do chọn đề tài")
    body_texts_ch1_1 = [
        "Trong thời đại cách mạng công nghiệp 4.0, dữ liệu được coi là nguồn tài nguyên quý giá nhất của nền kinh tế số. Đặc biệt trong lĩnh vực giao thông vận tải đô thị — nơi hàng triệu lượt di chuyển diễn ra mỗi ngày — việc thu thập, xử lý và phân tích dữ liệu một cách khoa học và có hệ thống là yếu tố then chốt để tối ưu hóa hoạt động, nâng cao chất lượng dịch vụ và hỗ trợ ra quyết định chiến lược.",
        "Transport for London (TfL) là một trong những tổ chức giao thông đô thị minh bạch và tiên tiến nhất thế giới. TfL cung cấp nguồn dữ liệu mở phong phú về lưu lượng hành khách, tuyến đường và cơ sở hạ tầng — tạo ra cơ hội lý tưởng để sinh viên ngành kỹ thuật dữ liệu thực hành trên bộ dữ liệu thực tế, quy mô lớn và mang tính học thuật cao.",
        "Ngoài ra, đại dịch COVID-19 đã gây ra những biến động lịch sử chưa từng có trong lịch sử giao thông công cộng toàn cầu. Dữ liệu TfL giai đoạn 2017–2021 bao phủ cả giai đoạn trước, trong và sau đại dịch, cung cấp một trường hợp nghiên cứu đặc biệt về sức chịu đựng và khả năng phục hồi của hệ thống giao thông đô thị.",
        "Xuất phát từ những lý do trên, Nhóm 14 đã lựa chọn đề tài \"Phân tích Hệ thống Giao thông Công cộng London (TfL)\" nhằm xây dựng một pipeline kỹ thuật dữ liệu hoàn chỉnh từ khâu thu thập, làm sạch, phân tích đến trực quan hóa, qua đó vận dụng toàn diện các kiến thức đã học trong môn Nhập môn Kỹ thuật Dữ liệu.",
    ]
    for text in body_texts_ch1_1:
        add_body_text(doc, text)
    
    add_heading_1(doc, "1.2", "Mục tiêu nghiên cứu")
    add_body_text(doc, "Dự án hướng đến các mục tiêu cụ thể sau:")
    
    goals = [
        "Xây dựng pipeline ETL hoàn chỉnh để thu thập và tích hợp dữ liệu từ ba nguồn dị biệt: file KML địa lý, CSV lưu lượng hành khách TfL và cơ sở dữ liệu NaPTAN.",
        "Áp dụng các kỹ thuật làm sạch và biến đổi dữ liệu nâng cao như chuẩn hóa tên ga bằng Regex, Spatial Join và Feature Engineering để tạo bộ dữ liệu chất lượng cao.",
        "Triển khai thuật toán phân cụm KMeans để phân loại 298 nhà ga thành các nhóm có đặc trưng riêng biệt về lưu lượng và vị trí địa lý.",
        "Phân tích định lượng tác động của đại dịch COVID-19 lên hệ thống giao thông công cộng London trong giai đoạn 2019–2021.",
        "Xây dựng hệ thống trực quan hóa tương tác dựa trên nền tảng Folium và Leaflet.js, cho phép khám phá dữ liệu đa chiều và xuất kết quả.",
        "Lưu trữ dữ liệu đã xử lý vào cơ sở dữ liệu SQLite và các định dạng chuẩn (CSV, Excel) để phục vụ các mục đích phân tích tiếp theo.",
    ]
    for i, goal in enumerate(goals, 1):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(f"({i}) {goal}")
        set_run_font(run, size=13)
    
    add_heading_1(doc, "1.3", "Phạm vi và đối tượng nghiên cứu")
    
    add_heading_2(doc, "1.3.1", "Phạm vi nghiên cứu")
    add_body_text(doc, "Về không gian: Dự án tập trung phân tích toàn bộ hệ thống nhà ga của Transport for London bao gồm tàu điện ngầm, tàu trên cao, DLR, tuyến Elizabeth và hệ thống xe điện Tram trên địa bàn thành phố London và vùng phụ cận.")
    add_body_text(doc, "Về thời gian: Dữ liệu được phân tích trong giai đoạn 5 năm từ 2017 đến 2021, bao phủ cả giai đoạn trước dịch (2017–2019), thời kỳ đỉnh dịch (2020) và giai đoạn phục hồi (2021).")
    add_body_text(doc, "Về kỹ thuật: Dự án giới hạn ở việc phân tích dữ liệu tĩnh (batch processing) và không xử lý dữ liệu thời gian thực (real-time streaming).")
    
    add_heading_2(doc, "1.3.2", "Đối tượng nghiên cứu")
    add_body_text(doc, "Đối tượng chính của nghiên cứu là 298 nhà ga thuộc hệ thống TfL đã qua kiểm tra và đủ điều kiện phân tích (có đầy đủ tọa độ địa lý, thông tin tuyến và dữ liệu lưu lượng hành khách năm 2021). Ngoài ra, nghiên cứu cũng xem xét các yếu tố địa lý (borough/quận) và chính sách ứng phó COVID-19 như các nhân tố bối cảnh quan trọng.")
    
    add_heading_1(doc, "1.4", "Ý nghĩa khoa học và thực tiễn")
    
    add_heading_2(doc, "1.4.1", "Ý nghĩa khoa học")
    add_body_text(doc, "Về mặt học thuật, dự án đóng góp vào cơ sở kiến thức về ứng dụng kỹ thuật dữ liệu trong phân tích giao thông đô thị. Pipeline ETL được thiết kế theo kiến trúc module hóa, có tính tái sử dụng cao và có thể dễ dàng mở rộng cho các bài toán phân tích dữ liệu giao thông tương tự ở các thành phố khác.")
    add_body_text(doc, "Việc kết hợp ba nguồn dữ liệu dị biệt (KML, CSV, NaPTAN) thông qua các thuật toán Regex matching và Nearest Neighbor là một đóng góp kỹ thuật thực tiễn trong việc xử lý vấn đề data fusion từ nhiều nguồn có chất lượng không đồng đều.")
    
    add_heading_2(doc, "1.4.2", "Ý nghĩa thực tiễn")
    add_body_text(doc, "Từ góc độ thực tiễn, kết quả phân tích cung cấp bức tranh toàn cảnh về sức khỏe của hệ thống giao thông công cộng London, giúp các nhà quản lý có thể xác định các nhà ga cần ưu tiên đầu tư, tái cơ cấu hoặc điều chỉnh tần suất phục vụ. Phân tích tác động COVID-19 cung cấp dữ liệu định lượng quan trọng để lập kế hoạch ứng phó cho các tình huống khủng hoảng tương tự trong tương lai.")
    add_body_text(doc, "Đối với sinh viên và cộng đồng học thuật, dự án cung cấp một case study hoàn chỉnh về quy trình kỹ thuật dữ liệu từ đầu đến cuối — từ thu thập dữ liệu thô đến sản phẩm cuối là bản đồ tương tác và báo cáo phân tích.")
    
    add_heading_1(doc, "1.5", "Cấu trúc báo cáo")
    add_body_text(doc, "Báo cáo được tổ chức thành 7 chương chính như sau:")
    
    structure = [
        ("Chương 1 — Mở đầu:", "Trình bày lý do chọn đề tài, mục tiêu, phạm vi nghiên cứu và ý nghĩa của dự án."),
        ("Chương 2 — Tổng quan về dữ liệu:", "Giới thiệu TfL, mô tả chi tiết ba nguồn dữ liệu đầu vào và đánh giá chất lượng dữ liệu thô."),
        ("Chương 3 — Thiết kế hệ thống:", "Trình bày kiến trúc pipeline ETL, thiết kế cơ sở dữ liệu và các công nghệ sử dụng."),
        ("Chương 4 — Triển khai và thực hiện:", "Mô tả chi tiết từng bước Extract, Transform và Load, thuật toán KMeans và hệ thống trực quan hóa."),
        ("Chương 5 — Kết quả và phân tích:", "Trình bày kết quả phân cụm, phân tích COVID-19 và các insight quan trọng."),
        ("Chương 6 — Đánh giá hệ thống:", "Phân tích ưu nhược điểm và đề xuất giải pháp cải thiện."),
        ("Chương 7 — Kết luận và kiến nghị:", "Tổng kết kết quả và định hướng phát triển tương lai."),
    ]
    
    for label, desc in structure:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        r1 = p.add_run(label + " ")
        set_run_font(r1, size=13, bold=True)
        r2 = p.add_run(desc)
        set_run_font(r2, size=13)
    
    add_page_break(doc)
    
    # ====================
    # CHƯƠNG 2: TỔNG QUAN VỀ DỮ LIỆU
    # ====================
    add_heading_chapter(doc, "2", "TỔNG QUAN VỀ DỮ LIỆU")
    
    add_heading_1(doc, "2.1", "Giới thiệu về Transport for London (TfL)")
    add_body_text(doc, "Transport for London (TfL) là cơ quan quản lý giao thông tích hợp trực thuộc Ủy ban Greater London, thành lập theo Đạo luật Greater London Authority năm 1999. TfL chịu trách nhiệm quản lý và vận hành toàn bộ mạng lưới giao thông công cộng của thủ đô London, bao gồm một trong những hệ thống tàu điện ngầm lâu đời nhất thế giới (London Underground, ra đời năm 1863) cùng nhiều phương thức giao thông hiện đại khác.")
    add_body_text(doc, "Hệ thống TfL bao gồm các phương thức giao thông chính sau: (1) London Underground (LU) với 11 tuyến và 272 nhà ga; (2) London Overground (LO) với mạng lưới tàu trên cao; (3) Docklands Light Railway (DLR); (4) Elizabeth Line (EL) — tuyến mới nhất; và (5) Tram tại Nam London. Mỗi năm trước đại dịch COVID-19, TfL phục vụ trên 3 tỷ lượt hành khách, tương đương khoảng 5 triệu lượt mỗi ngày làm việc.")
    add_body_text(doc, "TfL được biết đến với chính sách minh bạch dữ liệu (Open Data Policy), cung cấp API và bộ dữ liệu mở cho phép nhà nghiên cứu, lập trình viên và doanh nghiệp truy cập thông tin về lịch trình, lưu lượng hành khách và trạng thái hoạt động của hệ thống. Đây là nền tảng quan trọng cho việc phát triển các ứng dụng di động giao thông và nghiên cứu học thuật về giao thông đô thị.")
    
    add_heading_1(doc, "2.2", "Mô tả chi tiết các nguồn dữ liệu")
    add_body_text(doc, "Dự án sử dụng ba nguồn dữ liệu độc lập, mỗi nguồn cung cấp một khía cạnh khác nhau về hệ thống nhà ga TfL. Dưới đây là mô tả chi tiết từng nguồn:")
    
    # Table 2.1
    data_sources = [
        ("stations.kml / stations .kml", "KML (Google Earth)", "~102 KB", "Tọa độ GPS (lat/lon) và tên các nhà ga TfL"),
        ("TfL_stations.csv / stations.csv", "CSV (TfL Open Data)", "~79 KB", "Lưu lượng hành khách En/Ex 2017–2021, tuyến và mạng lưới"),
        ("Stops.csv", "CSV (NaPTAN Database)", "~101 MB", "Metadata điểm dừng: CommonName, Borough, StopType, Tọa độ"),
    ]
    add_table_with_caption(doc, "Bảng 2.1: Tổng quan ba nguồn dữ liệu đầu vào của dự án",
                           ["Tên file", "Định dạng", "Kích thước", "Nội dung chính"],
                           data_sources, col_widths=[4, 3, 2.5, 6.5])
    
    add_heading_2(doc, "2.2.1", "File KML — Dữ liệu tọa độ địa lý nhà ga")
    add_body_text(doc, "File KML (Keyhole Markup Language) là định dạng XML được phát triển bởi Google và chuẩn hóa theo ISO 19139. File stations.kml chứa tọa độ GPS (kinh độ, vĩ độ) và tên của các nhà ga TfL, được tổ chức theo cấu trúc Placemark — mỗi Placemark đại diện cho một nhà ga với thẻ <name> chứa tên ga và thẻ <coordinates> chứa tọa độ theo định dạng (lon,lat,alt).")
    add_body_text(doc, "Đây là nguồn dữ liệu địa lý chính của dự án, cung cấp nền tảng cho tất cả các thao tác liên quan đến vị trí không gian. Tuy nhiên, tên nhà ga trong file KML thường có hậu tố như \" Station\" và không khớp trực tiếp với tên trong các file CSV, đòi hỏi bước chuẩn hóa tên kỹ lưỡng trước khi ghép dữ liệu.")
    
    add_heading_2(doc, "2.2.2", "File TfL_stations.csv — Dữ liệu lưu lượng hành khách")
    add_body_text(doc, "File TfL_stations.csv là nguồn dữ liệu chính về lưu lượng hành khách, được lấy từ chương trình OpenData của TfL. File này chứa thống kê số lượt Vào/Ra (Entry/Exit — En/Ex) theo từng nhà ga cho các năm từ 2017 đến 2021, cùng thông tin về các tuyến đường phục vụ (cột LINES) và mạng lưới (cột NETWORK).")
    
    tfl_csv_cols = [
        ("Station", "Tên nhà ga (có thể khác với KML)"),
        ("En/Ex 2017", "Tổng lượt vào/ra năm 2017 (triệu lượt)"),
        ("En/Ex 2018", "Tổng lượt vào/ra năm 2018"),
        ("En/Ex 2019", "Tổng lượt vào/ra năm 2019"),
        ("En/Ex 2020", "Tổng lượt vào/ra năm 2020 (năm dịch)"),
        ("En/Ex 2021", "Tổng lượt vào/ra năm 2021"),
        ("LINES", "Danh sách tuyến đường phục vụ ga"),
        ("NETWORK", "Mạng lưới (Underground, Overground, DLR...)"),
        ("London Underground", "Có phục vụ LU không (Yes/No)"),
        ("Night Tube?", "Có hoạt động ban đêm không"),
    ]
    add_table_with_caption(doc, "Bảng 2.2: Cấu trúc dữ liệu file TfL_stations.csv",
                           ["Tên cột", "Mô tả"], tfl_csv_cols, col_widths=[4, 12])
    
    add_heading_2(doc, "2.2.3", "File Stops.csv — Cơ sở dữ liệu NaPTAN")
    add_body_text(doc, "File Stops.csv được lấy từ NaPTAN (National Public Transport Access Nodes) — cơ sở dữ liệu quốc gia của Anh về tất cả các điểm dừng giao thông công cộng. Với kích thước lên đến 101 MB, đây là nguồn dữ liệu lớn nhất trong dự án, chứa thông tin metadata phong phú về các điểm dừng, bao gồm tên thông dụng (CommonName), tọa độ GPS, loại điểm dừng (StopType) và thông tin địa lý về quận/thành phố (Borough).")
    
    stoptypes = [
        ("MET", "Metro (tàu điện ngầm)", "1"),
        ("RLY", "Railway (đường sắt quốc gia)", "2"),
        ("RSE", "Rail station entrance (cổng vào ga)", "3"),
        ("TMU", "Tram/metro stop (xe điện/metro)", "4"),
        ("DLR", "DLR stop (Docklands Light Railway)", "5"),
        ("PLT", "Platform (sân ga/bến đỗ)", "6"),
        ("RPL", "Rail platform (sân ga đường sắt)", "7"),
    ]
    add_table_with_caption(doc, "Bảng 2.3: Các loại StopType trong NaPTAN và mức độ ưu tiên",
                           ["StopType", "Mô tả", "Ưu tiên"], stoptypes,
                           col_widths=[2.5, 8, 2])
    
    add_heading_1(doc, "2.3", "Thách thức của dữ liệu thô")
    add_body_text(doc, "Việc kết hợp ba nguồn dữ liệu trên đặt ra nhiều thách thức kỹ thuật đáng kể:")
    
    challenges = [
        ("Không nhất quán về tên nhà ga:", "Cùng một nhà ga có thể có nhiều cách biểu diễn tên khác nhau giữa các nguồn, ví dụ: \"Bank Underground Station\" (KML), \"Bank\" (TfL CSV), \"Bank (Monument)\" (NaPTAN). Điều này đòi hỏi thuật toán chuẩn hóa mạnh mẽ dựa trên Regex."),
        ("Dữ liệu trùng lặp:", "File NaPTAN chứa nhiều bản ghi cho cùng một nhà ga (theo từng cổng vào, sân ga, hướng di chuyển). Ví dụ, ga Liverpool Street có thể xuất hiện hàng chục lần với StopType khác nhau."),
        ("Dữ liệu thiếu (Missing Values):", "Một số nhà ga không có đầy đủ thông tin lưu lượng hành khách cho tất cả các năm, đặc biệt do ảnh hưởng của COVID-19 và các thay đổi trong hệ thống báo cáo."),
        ("Định dạng số liệu không chuẩn:", "Số liệu lưu lượng hành khách trong TfL CSV chứa dấu phẩy ngăn cách (ví dụ: \"3,158,442\") và đôi khi bị bọc trong dấu ngoặc kép kép, yêu cầu xử lý CSV đặc biệt."),
        ("Tọa độ sai lệch:", "Một số trường hợp tọa độ trong file KML và NaPTAN có sai lệch nhỏ do phương pháp đo đạc khác nhau, đòi hỏi thuật toán chọn tọa độ tốt nhất dựa trên khoảng cách Euclidean."),
    ]
    
    for label, desc in challenges:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        r1 = p.add_run(f"• {label} ")
        set_run_font(r1, size=13, bold=True)
        r2 = p.add_run(desc)
        set_run_font(r2, size=13)
    
    add_heading_1(doc, "2.4", "Đánh giá chất lượng bộ dữ liệu")
    add_body_text(doc, "Sau khi tiến hành phân tích sơ bộ, bộ dữ liệu được đánh giá theo các tiêu chí chất lượng dữ liệu chuẩn (Data Quality Dimensions) như sau:")
    
    quality_data = [
        ("Tính đầy đủ (Completeness)", "Tốt", "~85% nhà ga có đủ 5 năm dữ liệu; ~15% thiếu một vài năm"),
        ("Tính chính xác (Accuracy)", "Tốt", "Dữ liệu nguồn gốc từ cơ quan nhà nước, độ tin cậy cao"),
        ("Tính nhất quán (Consistency)", "Trung bình", "Không nhất quán tên ga giữa 3 nguồn; cần chuẩn hóa"),
        ("Tính kịp thời (Timeliness)", "Tốt", "Dữ liệu cập nhật theo năm tài chính TfL"),
        ("Tính duy nhất (Uniqueness)", "Yếu (Stops.csv)", "NaPTAN có nhiều bản ghi trùng lặp cần dedup"),
    ]
    add_table_with_caption(doc, "Bảng 2.4: Đánh giá chất lượng dữ liệu theo các tiêu chí chuẩn",
                           ["Tiêu chí", "Đánh giá", "Ghi chú chi tiết"],
                           quality_data, col_widths=[4.5, 2.5, 9])
    
    add_page_break(doc)
    
    # ====================
    # CHƯƠNG 3: THIẾT KẾ HỆ THỐNG
    # ====================
    add_heading_chapter(doc, "3", "THIẾT KẾ HỆ THỐNG")
    
    add_heading_1(doc, "3.1", "Kiến trúc Pipeline ETL End-to-End")
    add_body_text(doc, "Pipeline ETL (Extract-Transform-Load) của dự án được thiết kế theo kiến trúc tuyến tính, tuần tự với 6 giai đoạn chính, mỗi giai đoạn được đóng gói thành các hàm Python độc lập và có thể kiểm thử riêng biệt (xem Hình 3.1). Kiến trúc này đảm bảo tính module hóa, dễ bảo trì và mở rộng.")
    
    # ETL architecture as text diagram
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run("[ Nguồn 1: KML ] ──┐")
    set_run_font(run, size=11)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run("[ Nguồn 2: TfL CSV ] ── EXTRACT ──► TRANSFORM ──► CLUSTERING ──► LOAD ──► VISUALIZATION")
    set_run_font(run, size=11)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run("[ Nguồn 3: NaPTAN ] ──┘")
    set_run_font(run, size=11)
    
    add_figure_caption(doc, "Hình 3.1: Kiến trúc tổng thể Pipeline ETL End-to-End")
    
    etl_stages = [
        ("Giai đoạn 1 — EXTRACT:", "load_kml_stations(), load_tfl_csv(), load_stops_csv(). Thu thập và đọc dữ liệu từ 3 nguồn, chuyển đổi sang DataFrame Pandas tiêu chuẩn."),
        ("Giai đoạn 2 — TRANSFORM:", "merge_sources(), clean_and_engineer(). Chuẩn hóa tên, ghép dữ liệu bằng station_key, tính các đặc trưng phân tích (covid_impact_pct, trend_slope...)."),
        ("Giai đoạn 3 — CLUSTERING:", "run_kmeans_clustering(). Phân cụm KMeans K=6 trên đặc trưng chuẩn hóa."),
        ("Giai đoạn 4 — LOAD:", "save_outputs(). Lưu ra CSV, Excel và SQLite."),
        ("Giai đoạn 5 — VISUALIZATION:", "create_folium_map(). Tạo bản đồ HTML tương tác."),
        ("Giai đoạn 6 — REPORT:", "generate_pdf_report(). Xuất báo cáo PDF tóm tắt."),
    ]
    
    for label, desc in etl_stages:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        r1 = p.add_run(f"• {label} ")
        set_run_font(r1, size=13, bold=True)
        r2 = p.add_run(desc)
        set_run_font(r2, size=13)
    
    add_heading_1(doc, "3.2", "Thiết kế cơ sở dữ liệu SQLite")
    add_body_text(doc, "Cơ sở dữ liệu SQLite london_tfl.db được thiết kế theo mô hình Star Schema đơn giản với hai bảng chính:")
    
    fact_stations = [
        ("station", "TEXT", "Tên nhà ga (chuẩn hóa)"),
        ("station_key", "TEXT", "Khóa chuẩn hóa để ghép dữ liệu"),
        ("lat", "REAL", "Vĩ độ GPS"),
        ("lon", "REAL", "Kinh độ GPS"),
        ("passengers_2017", "INTEGER", "Lượt hành khách năm 2017"),
        ("passengers_2018", "INTEGER", "Lượt hành khách năm 2018"),
        ("passengers_2019", "INTEGER", "Lượt hành khách năm 2019"),
        ("passengers_2020", "INTEGER", "Lượt hành khách năm 2020"),
        ("passengers_2021", "INTEGER", "Lượt hành khách năm 2021"),
        ("num_lines", "INTEGER", "Số tuyến phục vụ nhà ga"),
        ("avg_passengers", "REAL", "Trung bình hành khách 5 năm"),
        ("covid_impact_pct", "REAL", "Tác động COVID-19 (% thay đổi 2019→2020)"),
        ("recovery_rate_pct", "REAL", "Tỷ lệ phục hồi (% thay đổi 2020→2021)"),
        ("trend_slope", "REAL", "Hệ số xu hướng (LinearRegression)"),
        ("trend_category", "TEXT", "Phân loại xu hướng (Tăng mạnh/Giảm mạnh...)"),
        ("cluster_id", "INTEGER", "ID cụm KMeans (0–5)"),
        ("cluster_name", "TEXT", "Tên cụm (Siêu trung tâm/Ga lớn...)"),
        ("borough", "TEXT", "Quận/Borough của nhà ga"),
    ]
    add_table_with_caption(doc, "Bảng 3.2: Cấu trúc bảng fact_stations trong SQLite",
                           ["Tên cột", "Kiểu dữ liệu", "Mô tả"],
                           fact_stations, col_widths=[4, 2.5, 9.5])
    
    add_heading_1(doc, "3.3", "Công nghệ, thư viện và công cụ sử dụng")
    
    tech_stack = [
        ("Python 3.10+", "Ngôn ngữ lập trình chính của dự án"),
        ("Pandas", "Xử lý và phân tích dữ liệu dạng bảng (DataFrame)"),
        ("NumPy", "Tính toán số học và mảng đa chiều"),
        ("Scikit-learn", "Triển khai KMeans clustering và StandardScaler"),
        ("Folium / Leaflet.js", "Tạo bản đồ tương tác HTML"),
        ("SQLite3 (stdlib)", "Lưu trữ dữ liệu trong cơ sở dữ liệu nhẹ"),
        ("OpenPyXL", "Xuất dữ liệu ra định dạng Excel (.xlsx)"),
        ("xml.etree.ElementTree", "Đọc và phân tích file KML (XML)"),
        ("ReportLab", "Tạo file PDF báo cáo tóm tắt"),
        ("Regex (re)", "Chuẩn hóa và làm sạch chuỗi tên ga"),
        ("pyngrok / localtunnel", "Tạo public URL để chia sẻ bản đồ qua internet"),
        ("http.server (stdlib)", "Web server tích hợp để phục vụ file HTML"),
    ]
    add_table_with_caption(doc, "Bảng 3.1: Các thư viện và công cụ sử dụng trong dự án",
                           ["Thư viện / Công cụ", "Vai trò và chức năng"],
                           tech_stack, col_widths=[4.5, 11.5])
    
    add_page_break(doc)
    
    # ====================
    # CHƯƠNG 4: TRIỂN KHAI VÀ THỰC HIỆN
    # ====================
    add_heading_chapter(doc, "4", "TRIỂN KHAI VÀ THỰC HIỆN")
    
    add_heading_1(doc, "4.1", "Extract — Thu thập dữ liệu từ nhiều nguồn")
    
    add_heading_2(doc, "4.1.1", "Đọc file KML — Hàm load_kml_stations()")
    add_body_text(doc, "Hàm load_kml_stations() sử dụng thư viện xml.etree.ElementTree để đọc và phân tích cú pháp file KML. Một thách thức kỹ thuật quan trọng là file KML thực tế có namespace XML (ví dụ {http://www.opengis.net/kml/2.2}), khiến việc tìm kiếm phần tử bằng XPath thông thường thất bại. Giải pháp được áp dụng là loại bỏ namespace khỏi tên thẻ trước khi xử lý (strip namespace).")
    add_body_text(doc, "Quá trình xử lý: Với mỗi phần tử <Placemark>, hàm trích xuất tên ga từ <name> và tọa độ từ <coordinates> (định dạng lon,lat,alt). Sau khi parse, tên ga được làm sạch bằng cách loại bỏ hậu tố \" Station\" và chuẩn hóa khoảng trắng. Kết quả là DataFrame với 4 cột: station, station_key, lat, lon.")
    
    add_heading_2(doc, "4.1.2", "Đọc file TfL CSV — Hàm load_tfl_csv()")
    add_body_text(doc, "File TfL_stations.csv có một vấn đề đặc biệt: mỗi dòng dữ liệu bị bọc trong dấu ngoặc kép như một chuỗi đơn, thay vì được phân tách thành các trường riêng biệt. Hàm load_tfl_csv() xử lý vấn đề này bằng cách đọc file dưới dạng văn bản thuần túy, sau đó áp dụng csv.reader hai lần để 'un-escape' chuỗi lồng nhau.")
    add_body_text(doc, "Sau khi parse thành công, các cột lưu lượng hành khách (En/Ex YYYY) được chuyển đổi từ chuỗi có dấu phẩy sang số nguyên. Các nhà ga trùng station_key (do cùng một nhà ga phục vụ nhiều tuyến) được gộp lại bằng groupby với hàm tổng hợp phù hợp: sum() cho số liệu hành khách, combine_lines() cho danh sách tuyến.")
    
    add_heading_2(doc, "4.1.3", "Đọc file Stops.csv — Hàm load_stops_csv()")
    add_body_text(doc, "File Stops.csv với kích thước 101MB được đọc bằng pd.read_csv() với tùy chọn low_memory=False để tránh lỗi kiểu dữ liệu. Hàm xử lý thông tin địa lý (Borough) bằng cách ưu tiên lấy từ các cột Town, ParentLocalityName và LocalityName theo thứ tự. Mức độ ưu tiên StopType được gán theo bảng STOPTYPE_PRIORITY, với MET (tàu điện ngầm) được ưu tiên cao nhất.")
    
    add_heading_1(doc, "4.2", "Transform — Làm sạch và biến đổi dữ liệu")
    
    add_heading_2(doc, "4.2.1", "Chuẩn hóa tên ga bằng Regex")
    add_body_text(doc, "Hàm normalize_station_name() là trọng tâm của giai đoạn Transform. Hàm này thực hiện 5 bước chuẩn hóa tuần tự:")
    
    normalize_steps = [
        "Chuyển tất cả về chữ thường (lowercase) và loại bỏ khoảng trắng đầu/cuối.",
        "Thay thế ký hiệu đặc biệt: '&' → 'and', 'St.' → 'st'.",
        "Loại bỏ hậu tố phương thức vận chuyển theo vòng lặp: ' underground station', ' overground station', ' dlr station', ' elizabeth line station', ' tram stop', ' station' và các hậu tố mã tuyến như ' lu', ' lo', ' dlr', ' nr'.",
        "Xóa ký tự đặc biệt không phải chữ-số-dấu nháy bằng regex [^\\w\\s'].",
        "Áp dụng các alias thủ công cho các trường hợp đặc biệt: 'bank'/'monument' → 'bank and monument', 'heathrow terminals 123' → 'heathrow terminals 1 2 3'.",
    ]
    for idx, step in enumerate(normalize_steps):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        letter = chr(97 + idx)
        run = p.add_run(f"({letter}) {step}")
        set_run_font(run, size=13)
    
    regex_examples = [
        ("\"Bank Underground Station\"", "\"bank and monument\""),
        ("\"Liverpool Street NR\"", "\"liverpool street\""),
        ("\"Heathrow Terminals 1 2 3\"", "\"heathrow terminals 1 2 3\""),
        ("\"Hammersmith (D & P)\"", "\"hammersmith\""),
        ("\"King's Cross St. Pancras\"", "\"king's cross st pancras\""),
    ]
    add_table_with_caption(doc, "Bảng 4.1: Kết quả Regex chuẩn hóa tên ga — ví dụ minh họa",
                           ["Tên gốc", "Tên sau chuẩn hóa"],
                           regex_examples, col_widths=[7, 9])
    
    add_heading_2(doc, "4.2.2", "Ghép dữ liệu và Nearest Neighbor Matching")
    add_body_text(doc, "Hàm merge_sources() thực hiện ghép ba nguồn dữ liệu theo station_key bằng pd.merge() với phương thức LEFT JOIN để giữ toàn bộ nhà ga từ KML. Sau khi ghép KML với TfL CSV, hàm nearest_stop_match() tiếp tục ghép với NaPTAN để enrich thông tin Borough.")
    add_body_text(doc, "Vì một tên ga có thể xuất hiện nhiều lần trong NaPTAN (theo từng sân ga, cổng vào...), thuật toán chọn bản ghi tốt nhất dựa trên hai tiêu chí theo thứ tự ưu tiên: (1) StopType được ưu tiên theo bảng STOPTYPE_PRIORITY (MET > RLY > RSE > TMU > DLR > PLT > RPL); (2) Khoảng cách Euclidean nhỏ nhất so với tọa độ KML.")
    add_body_text(doc, "Thuật toán tính khoảng cách Euclidean dựa trên tọa độ GPS (kinh độ, vĩ độ) bằng công thức: d = √((lat₁ - lat₂)² + (lon₁ - lon₂)²). Do phạm vi địa lý của toàn mạng lưới London tương đối hẹp (trong bán kính khoảng 25km), việc tính toán Euclidean trên mặt phẳng hai chiều thay cho công thức Haversine có sai số cực kỳ nhỏ (dưới 0.3%). Sai số này hoàn toàn không làm thay đổi thứ tự ga gần nhất nhưng tăng tốc độ tính toán lên gấp nhiều lần, giúp tối ưu hóa đáng kể tài nguyên hệ thống.")
    
    merge_results = [
        ("Nguồn KML", "Tổng số nhà ga đọc được từ file KML"),
        ("Ghép TfL CSV", "Số nhà ga khớp được với TfL CSV qua station_key"),
        ("Ghép NaPTAN", "Số nhà ga có thông tin Borough từ Stops.csv"),
        ("Sau lọc cuối", "Số nhà ga đủ điều kiện phân tích (có lat/lon/passengers/lines)"),
    ]
    add_table_with_caption(doc, "Bảng 4.2: Thống kê kết quả quá trình ghép dữ liệu",
                           ["Bước", "Mô tả"], merge_results, col_widths=[4, 12])
    
    add_heading_2(doc, "4.2.3", "Feature Engineering")
    add_body_text(doc, "Hàm clean_and_engineer() tạo ra các đặc trưng phân tích quan trọng:")
    
    features = [
        ("num_lines", "Số tuyến phục vụ nhà ga, tính bằng cách đếm dấu phẩy trong cột LINES + 1."),
        ("avg_passengers", "Trung bình lưu lượng hành khách qua 5 năm (2017–2021)."),
        ("covid_impact_pct", "Phần trăm thay đổi lưu lượng 2019→2020: (p2020 - p2019) / p2019 × 100."),
        ("recovery_rate_pct", "Phần trăm thay đổi lưu lượng 2020→2021: (p2021 - p2020) / p2020 × 100."),
        ("trend_slope", "Hệ số hồi quy tuyến tính (LinearRegression) của lưu lượng theo năm 2017–2021."),
        ("trend_category", "Phân loại xu hướng dựa trên trend_slope: Tăng mạnh (>150K), Tăng nhẹ, Ổn định, Giảm nhẹ, Giảm mạnh (<-150K)."),
    ]
    for idx, (fname, fdesc) in enumerate(features):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        letter = chr(97 + idx)
        r1 = p.add_run(f"({letter}) {fname}: ")
        set_run_font(r1, size=13, bold=True)
        r2 = p.add_run(fdesc)
        set_run_font(r2, size=13)
    
    add_heading_1(doc, "4.3", "Phân cụm KMeans")
    
    add_heading_2(doc, "4.3.1", "Chọn số cụm K=6")
    add_body_text(doc, "Sau khi khảo sát bằng phương pháp Elbow (xem Hình 4.3) và xem xét ý nghĩa kinh doanh của từng nhóm, nhóm nghiên cứu chọn K=6 là số cụm tối ưu. K=6 cho phép phân biệt rõ ràng 6 cấp độ hoạt động khác nhau của nhà ga từ siêu trung tâm đến ga rất ít khách, đồng thời đủ chi tiết để cung cấp thông tin hữu ích mà không quá phức tạp để diễn giải.")
    add_body_text(doc, "Về số liệu toán học thực nghiệm, khi khảo sát các giá trị K chạy từ 2 đến 10, nếu chỉ chọn theo Silhouette Score cao nhất thì K=2 là tối ưu nhất (Silhouette = 0.4450). Tuy nhiên, phân cụm này quá thô khi dồn hầu hết ga vào 1 cụm. Khi tăng K lên 6, mô hình đạt điểm cân bằng lý tưởng: tổng bình phương khoảng cách nội cụm (Inertia) giảm mạnh từ 864.8 (K=2) xuống 364.3 (K=6) — tương ứng giảm 57.8% độ phân tán của dữ liệu. Đồng thời, con số K=6 cũng khớp hoàn chỉnh với cấu trúc phân chia 6 vùng Zone (Zone 1 đến Zone 6) phục vụ giá vé và phân vùng địa lý thực tế của Transport for London.")
    
    add_heading_2(doc, "4.3.2", "Đặc trưng đầu vào và quy trình phân cụm")
    
    cluster_features = [
        ("passengers_2021", "Lưu lượng hành khách năm 2021 — thước đo chính về tầm quan trọng hiện tại"),
        ("num_lines", "Số tuyến phục vụ — thể hiện tính kết nối và vai trò trong mạng lưới"),
        ("lat", "Vĩ độ GPS — vị trí địa lý Bắc/Nam"),
        ("lon", "Kinh độ GPS — vị trí địa lý Đông/Tây"),
    ]
    add_table_with_caption(doc, "Bảng 4.3: Các đặc trưng đầu vào cho mô hình KMeans",
                           ["Đặc trưng", "Ý nghĩa và lý do chọn"],
                           cluster_features, col_widths=[4, 12])
    
    add_body_text(doc, "Quy trình phân cụm: (1) Chuẩn hóa 4 đặc trưng về phân phối chuẩn bằng StandardScaler để đảm bảo mỗi đặc trưng có đóng góp đồng đều; (2) Chạy KMeans với n_clusters=6, random_state=42 để đảm bảo tái hiện kết quả (reproducibility), n_init=10 để tránh local minima; (3) Đặt tên cụm theo thứ tự lưu lượng hành khách trung bình giảm dần.")
    
    add_heading_1(doc, "4.4", "Trực quan hóa tương tác với Folium")
    add_body_text(doc, "Hàm create_folium_map() tạo ra file london_tfl_map.html — một ứng dụng web tương tác hoàn chỉnh dựa trên thư viện Leaflet.js, MarkerCluster và Bootstrap 5. Bản đồ cung cấp đầy đủ các tính năng sau:")
    
    folium_features = [
        "Hiển thị 298 nhà ga dưới dạng CircleMarker với màu sắc phân biệt theo 6 cụm.",
        "Kích thước marker tỷ lệ với lưu lượng hành khách theo năm được chọn.",
        "MarkerCluster: Tự động nhóm các marker khi zoom ra để tránh chồng chéo.",
        "Time Slider: Thanh trượt chọn năm phân tích (2017–2021) để xem sự thay đổi theo thời gian.",
        "Bộ lọc cụm: Chips toggle để bật/tắt từng nhóm nhà ga.",
        "Bộ lọc Borough: Dropdown lọc theo quận/Borough.",
        "Tìm kiếm nhà ga: Tìm kiếm real-time theo tên, tuyến hoặc Borough.",
        "Heatmap: Bản đồ nhiệt thể hiện mật độ hành khách.",
        "Popup chi tiết: Click vào marker để xem thông tin đầy đủ kèm biểu đồ sparkline.",
        "Xuất CSV: Nút export dữ liệu đang lọc ra file CSV.",
        "Dark mode: Hỗ trợ chế độ bản đồ tối.",
    ]
    for idx, feat in enumerate(folium_features):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        letter = chr(97 + idx)
        run = p.add_run(f"({letter}) {feat}")
        set_run_font(run, size=13)
    
    add_heading_1(doc, "4.5", "Lưu trữ và xuất dữ liệu")
    add_body_text(doc, "Hàm save_outputs() lưu kết quả vào thư mục outputs/ với bốn định dạng:")
    add_body_text(doc, "CSV (london_tfl_cleaned.csv): Định dạng phổ biến nhất, dễ mở bằng Excel hoặc đọc bằng Python/R. Phù hợp để chia sẻ và import vào các công cụ BI.")
    add_body_text(doc, "Excel (london_tfl_results.xlsx): Chứa hai sheet: All_Stations với toàn bộ dữ liệu nhà ga và Cluster_Summary với bảng tổng hợp theo cụm. Sử dụng thư viện OpenPyXL.")
    add_body_text(doc, "SQLite (london_tfl.db): Cơ sở dữ liệu quan hệ với hai bảng: fact_stations (toàn bộ dữ liệu nhà ga) và dim_clusters (tổng hợp theo cụm). Tên cột được làm sạch để tránh xung đột với reserved keywords của SQLite.")
    add_body_text(doc, "Text Report (analysis_summary.txt): Báo cáo tóm tắt dạng văn bản với các số liệu thống kê chính, phù hợp để in nhanh hoặc gửi email.")
    
    add_page_break(doc)
    
    # ====================
    # CHƯƠNG 5: KẾT QUẢ VÀ PHÂN TÍCH
    # ====================
    add_heading_chapter(doc, "5", "KẾT QUẢ VÀ PHÂN TÍCH")
    
    add_heading_1(doc, "5.1", "Kết quả phân cụm KMeans (6 cụm)")
    add_body_text(doc, "Sau khi chạy thuật toán KMeans với K=6 trên 298 nhà ga, kết quả phân cụm được tổng hợp trong Bảng 5.1. Mỗi cụm được đặt tên theo mức độ lưu lượng hành khách trung bình năm 2021, từ \"Siêu trung tâm\" (đông nhất) đến \"Ga rất ít khách\" (ít nhất).")
    
    cluster_results = [
        ("Siêu trung tâm", "8", "37,315,886", "−73.70%", "+71.82%", "Stratford, Liverpool Street, King's Cross St. Pancras, Victoria, Oxford Circus, Bank, Waterloo, Paddington"),
        ("Ga lớn", "42", "11,268,126", "−63.47%", "+39.17%", "London Bridge, Canary Wharf, Highbury & Islington, Brixton, Whitechapel, ..."),
        ("Ga trung bình", "78", "3,136,520", "−54.91%", "+0.44%", "Angel, Aldgate, Bethnal Green, Canning Town, ..."),
        ("Ga nhỏ", "80", "2,833,787", "−57.42%", "+27.60%", "Các ga ngoại ô nội đô, bến đỗ dân cư"),
        ("Ga ít khách", "37", "2,215,412", "−49.69%", "−10.64%", "Các ga vùng ven nội thành, lưu lượng thấp"),
        ("Ga rất ít khách", "53", "1,630,990", "−41.42%", "−20.38%", "Các ga ngoại ô xa, Zone 5-6 hoặc Tram stop"),
    ]
    add_table_with_caption(doc, "Bảng 5.1: Tổng hợp kết quả phân cụm KMeans 6 cụm",
                           ["Cụm", "Số ga", "HK TB 2021", "COVID Impact", "Phục hồi", "Ví dụ ga tiêu biểu"],
                           cluster_results, col_widths=[3, 1.5, 3, 2.5, 2, 4])
    
    add_body_text(doc, "Phân tích kết quả cho thấy sự phân cấp rõ ràng về lưu lượng hành khách: cụm Siêu trung tâm có lưu lượng trung bình gấp 23 lần cụm Ga rất ít khách. Đặc biệt, các ga thuộc cụm Siêu trung tâm bị ảnh hưởng nặng nề nhất bởi COVID-19 (−73.70%) do phụ thuộc nhiều vào hành khách công sở và du lịch quốc tế — hai nhóm đối tượng bị hạn chế nhất trong đại dịch.")
    add_body_text(doc, "Về phân bố địa lý, cụm Siêu trung tâm tập trung chủ yếu ở khu vực Zone 1 (trung tâm London) và một số điểm trung chuyển lớn ở Zone 2. Các cụm Ga rất ít khách phân bố rải rác ở vùng ngoại ô, Zone 4–6, nơi mạng lưới giao thông công cộng kém phát triển hơn và người dân có xu hướng sử dụng phương tiện cá nhân.")
    
    add_heading_1(doc, "5.2", "Phân tích tác động COVID-19")
    add_body_text(doc, "Kết quả phân tích tác động COVID-19 từ dữ liệu TfL giai đoạn 2019–2021 cung cấp bằng chứng định lượng rõ ràng về ảnh hưởng của đại dịch:")
    
    covid_data = [
        ("2019 (trước dịch)", "3,283,909,448", "—", "Mức cao nhất trong giai đoạn nghiên cứu"),
        ("2020 (năm dịch bùng phát)", "1,208,311,783", "−63.21%", "Giảm kỷ lục do phong tỏa từ tháng 3/2020"),
        ("2021 (phục hồi)", "1,411,552,644", "+16.82%", "Phục hồi dần nhưng vẫn ở mức 43% so với 2019"),
    ]
    add_table_with_caption(doc, "Bảng 5.2: Tác động COVID-19 lên tổng lưu lượng hành khách TfL",
                           ["Năm", "Tổng lượt hành khách", "Thay đổi YoY", "Ghi chú"],
                           covid_data, col_widths=[3.5, 4.5, 3, 5])
    
    add_body_text(doc, "Phân tích sâu theo xu hướng 5 năm cho thấy 287/298 ga (96.3%) đang trong xu hướng \"Giảm mạnh\", phản ánh tác động nghiêm trọng của COVID-19. Chỉ có 9 ga có xu hướng \"Giảm nhẹ\", 1 ga \"Tăng mạnh\" (Elizabeth Line mới) và 1 ga \"Ổn định\".")
    add_body_text(doc, "Điểm đáng chú ý là các ga ở cụm Ga rất ít khách bị ảnh hưởng ít nhất (−41.42%) trong khi các ga Siêu trung tâm bị ảnh hưởng nặng nhất (−73.70%). Điều này có thể được giải thích bởi thực tế rằng các ga ít khách phục vụ chủ yếu hành khách thiết yếu (công nhân, y tế, thương mại địa phương) trong khi các ga trung tâm phụ thuộc nhiều vào người đi làm văn phòng và khách du lịch.")
    
    add_heading_1(doc, "5.3", "Các insight quan trọng và Top ga đông khách")
    
    top10_stations = [
        ("1", "Stratford", "Siêu trung tâm", "63,439,192", "−54.6%", "Ga kết nối 5 tuyến lớn, bến đỗ đông khách nhất sau Olympic"),
        ("2", "Liverpool Street", "Siêu trung tâm", "43,057,801", "−71.5%", "Hub kết nối khu tài chính City và Đông London"),
        ("3", "King's Cross St. Pancras", "Siêu trung tâm", "36,734,085", "−78.7%", "Ga liên thông 6 tuyến metro + đường sắt nội địa/quốc tế"),
        ("4", "Victoria", "Siêu trung tâm", "33,480,926", "−73.1%", "Cửa ngõ kết nối phía Nam và sân bay Gatwick"),
        ("5", "Oxford Circus", "Siêu trung tâm", "32,863,152", "−81.3%", "Trục mua sắm sầm uất bậc nhất trung tâm London"),
        ("6", "London Bridge", "Ga lớn", "30,857,082", "−66.7%", "Hub giao thông lớn kết nối phía Nam qua sông Thames"),
        ("7", "Bank", "Siêu trung tâm", "30,132,853", "−83.1%", "Tâm điểm khu tài chính City, kết nối 6 tuyến chạy"),
        ("8", "Waterloo", "Siêu trung tâm", "29,868,623", "−80.0%", "Từng là ga đông khách nhất, chịu ảnh hưởng nặng do dịch"),
        ("9", "Paddington", "Siêu trung tâm", "28,950,455", "−67.4%", "Cửa ngõ kết nối phía Tây và tuyến Elizabeth Line mới"),
        ("10", "Canary Wharf", "Ga lớn", "24,923,667", "−73.4%", "Trọng điểm khu tài chính mới phía Đông London"),
    ]
    add_table_with_caption(doc, "Bảng 5.3: Top 10 nhà ga đông khách nhất năm 2021",
                           ["STT", "Tên ga", "Cụm", "HK 2021", "COVID Impact", "Đặc điểm"],
                           top10_stations, col_widths=[1, 4.5, 3, 3, 2.5, 2])
    
    add_body_text(doc, "Một số insight quan trọng từ quá trình phân tích:")
    
    insights = [
        "Hiệu ứng COVID phân tầng: Các ga phục vụ khách du lịch và công sở bị ảnh hưởng nặng hơn (−70% đến −76%) so với các ga phục vụ nhu cầu thiết yếu (−40% đến −50%). Điều này gợi ý cần có chiến lược phục hồi khác nhau cho từng phân khúc nhà ga.",
        "Tập trung lưu lượng cao: 8 ga thuộc cụm Siêu trung tâm (2.7% tổng số ga) chiếm khoảng 21% tổng lưu lượng hành khách toàn mạng lưới năm 2021, thể hiện sự tập trung cao độ về nhu cầu giao thông.",
        "Tiềm năng phục hồi: Năm 2021 tổng lưu lượng đạt 43% so với mức 2019 — nhanh hơn dự kiến ban đầu, cho thấy nhu cầu sử dụng giao thông công cộng được dồn nén sẽ bùng nổ khi hạn chế được dỡ bỏ.",
        "Mạng lưới đa tuyến: Các ga phục vụ nhiều tuyến (5+ tuyến) có xu hướng phục hồi nhanh hơn so với ga đơn tuyến, nhờ tính linh hoạt và nhiều hướng tiếp cận hơn cho hành khách.",
    ]
    for idx, ins in enumerate(insights):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        letter = chr(97 + idx)
        run = p.add_run(f"({letter}) {ins}")
        set_run_font(run, size=13)
    
    add_heading_1(doc, "5.4", "Đánh giá hiệu suất mô hình")
    add_body_text(doc, "Mô hình KMeans được đánh giá bằng hai chỉ số định lượng:")
    add_body_text(doc, "Inertia (Within-cluster Sum of Squares): Đo độ chặt chẽ nội cụm — giá trị càng nhỏ càng tốt. Kết quả cho thấy với K=6, inertia đạt mức tối ưu theo biểu đồ Elbow, sau đó cải thiện không đáng kể khi tăng K.")
    add_body_text(doc, "Silhouette Score: Đo tính phân tách giữa các cụm (−1 đến 1, càng gần 1 càng tốt). Với K=6, silhouette score đạt khoảng 0.42 — mức \"cụm hợp lý\" theo thang đánh giá chuẩn.")
    add_body_text(doc, "Ngoài ra, tính giải thích được (Interpretability) của 6 cụm rất cao — mỗi cụm có đặc trưng rõ ràng về lưu lượng hành khách và vị trí địa lý, dễ dàng đặt tên và truyền đạt ý nghĩa đến các bên liên quan không chuyên kỹ thuật.")
    
    add_page_break(doc)
    
    # ====================
    # CHƯƠNG 6: ĐÁNH GIÁ HỆ THỐNG
    # ====================
    add_heading_chapter(doc, "6", "ĐÁNH GIÁ HỆ THỐNG")
    
    add_heading_1(doc, "6.1", "Ưu điểm của hệ thống")
    
    pros = [
        ("Thiết kế module hóa cao:", "Mỗi bước ETL được tách thành hàm độc lập (load_kml_stations, load_tfl_csv, merge_sources, run_kmeans_clustering...), dễ kiểm thử đơn vị (unit testing) và thay thế từng component."),
        ("Xử lý lỗi chủ động:", "Pipeline tích hợp nhiều cơ chế phát hiện và xử lý lỗi như kiểm tra file tồn tại (find_file), xử lý NaN trong chia phần (BUG FIX #3), cảnh báo khi số cụm không đủ dữ liệu (BUG FIX #5)."),
        ("Tự động hóa hoàn toàn:", "Toàn bộ pipeline từ đọc dữ liệu thô đến xuất kết quả và tạo bản đồ chạy tự động bằng lệnh python final_project.py, không cần can thiệp thủ công."),
        ("Trực quan hóa phong phú:", "Bản đồ Folium cung cấp trải nghiệm người dùng premium với đầy đủ tính năng lọc, tìm kiếm, heatmap và biểu đồ thống kê tương tác."),
        ("Đa định dạng xuất:", "Kết quả được lưu đồng thời ở 4 định dạng (CSV, Excel, SQLite, Text) đáp ứng nhu cầu của nhiều đối tượng người dùng khác nhau."),
        ("Tương thích đa nền tảng:", "Code chạy được trên cả Windows, macOS, Linux và Google Colab, với xử lý UTF-8 encoding đặc biệt cho Windows console."),
    ]
    
    for idx, (label, desc) in enumerate(pros):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        letter = chr(97 + idx)
        r1 = p.add_run(f"({letter}) {label} ")
        set_run_font(r1, size=13, bold=True, color=(0, 128, 0))
        r2 = p.add_run(desc)
        set_run_font(r2, size=13)
    
    add_heading_1(doc, "6.2", "Xử lý ngoại lệ dữ liệu (Quality Control & Bug Fixes)")
    add_body_text(doc, "Một hệ thống Kỹ thuật dữ liệu chuyên nghiệp đòi hỏi khả năng hoạt động ổn định và tự động xử lý các tình huống dị thường của dữ liệu. Trong quá trình xây dựng dự án, nhóm nghiên cứu đã tích hợp 5 cơ chế phòng ngừa lỗi và kiểm soát chất lượng dữ liệu:")
    
    qc_mechanisms = [
        ("Kiểm soát tính toàn vẹn khi ghép dữ liệu:", "Sử dụng df.isna().all() trên các trường quan trọng thay vì chỉ dùng df.empty để phát hiện lỗi ghép nối rỗng khi kết quả trả về các dòng toàn giá trị NaN ở phép ghép LEFT JOIN."),
        ("Bộ lọc CSV tùy chỉnh thông minh:", "Giải quyết lỗi double-parsing khi gặp các trường số chứa dấu phẩy ngăn cách hàng nghìn bị bọc trong dấu ngoặc kép. Parser sẽ tự động bóc lớp ngoặc kép ngoài cùng và phân tách lại các cột con bằng dấu phẩy."),
        ("Phòng ngừa lỗi chia cho 0 (ZeroDivisionError):", "Khi tính toán tỷ lệ sụt giảm do dịch bệnh (covid_impact_pct) hoặc phục hồi (recovery_rate_pct), hệ thống sử dụng np.where() kiểm tra mẫu số. Nếu mẫu số bằng 0 hoặc NaN, giá trị trả về được gán là NaN thay vì crash hệ thống."),
        ("Tính toán xu hướng linh hoạt:", "Trong mô hình Linear Regression phân tích xu hướng hành khách qua 5 năm, hệ thống sử dụng mặt nạ valid_mask để loại bỏ các điểm dữ liệu khuyết thiếu. Yêu cầu tối thiểu có 2 điểm dữ liệu thực tế (2 năm) để tiến hành fit đường thẳng hồi quy, ngăn ngừa lỗi Singular Matrix khi tính toán ma trận."),
        ("Tự động thích ứng số cụm KMeans:", "Hàm phân cụm KMeans sử dụng logic min(n_clusters, len(df)) đề phòng trường hợp tập dữ liệu lọc đầu vào có số lượng nhà ga nhỏ hơn số cụm yêu cầu (K=6), tránh lỗi phân cụm rỗng.")
    ]
    for idx, (title, desc) in enumerate(qc_mechanisms):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        letter = chr(97 + idx)
        r1 = p.add_run(f"({letter}) {title} ")
        set_run_font(r1, size=13, bold=True)
        r2 = p.add_run(desc)
        set_run_font(r2, size=13)
        
    add_heading_1(doc, "6.3", "Hạn chế và khó khăn gặp phải")
    
    cons = [
        ("Hiệu năng xử lý:", "File Stops.csv với 101MB phải load toàn bộ vào RAM, gây chậm trên máy tính yếu. Quá trình ghép dữ liệu theo station_key dùng pd.merge() chưa tối ưu về bộ nhớ."),
        ("Chỉ xử lý batch:", "Pipeline chỉ phân tích dữ liệu lịch sử tĩnh, không hỗ trợ real-time streaming. Không có cơ chế cập nhật tăng dần (incremental update) khi có dữ liệu mới."),
        ("Phụ thuộc chất lượng khớp tên:", "Thuật toán chuẩn hóa Regex có thể gặp khó khăn với các tên ga rất khác nhau giữa hai nguồn (ví dụ: địa danh bằng tiếng Anh cổ, từ viết tắt đặc biệt)."),
        ("Thiếu unit test:", "Code không có bộ kiểm thử tự động, khó phát hiện regression khi thay đổi code."),
        ("Bản đồ phụ thuộc internet:", "File HTML yêu cầu kết nối internet để load thư viện Leaflet.js và CartoDB tiles từ CDN."),
    ]
    
    for idx, (label, desc) in enumerate(cons):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        letter = chr(97 + idx)
        r1 = p.add_run(f"({letter}) {label} ")
        set_run_font(r1, size=13, bold=True, color=(192, 0, 0))
        r2 = p.add_run(desc)
        set_run_font(r2, size=13)
    
    add_heading_1(doc, "6.4", "Giải pháp khắc phục")
    
    solutions = [
        ("Tối ưu hiệu năng:", "Sử dụng Dask hoặc PySpark thay cho Pandas để xử lý file lớn, áp dụng lazy loading và chunked processing cho Stops.csv."),
        ("Real-time streaming:", "Tích hợp Apache Kafka hoặc AWS Kinesis để xử lý dữ liệu theo luồng thời gian thực từ API TfL."),
        ("Cải thiện khớp tên:", "Tích hợp thuật toán Fuzzy Matching (thư viện rapidfuzz) để tăng tỷ lệ khớp tên ga, đặc biệt cho các trường hợp gõ sai chính tả."),
        ("Bổ sung unit test:", "Viết test suite sử dụng pytest với fixture dữ liệu mẫu để đảm bảo mỗi hàm hoạt động đúng."),
        ("Bundle offline:", "Sử dụng Folium bundled mode hoặc webpack để đóng gói thư viện JavaScript vào file HTML, cho phép dùng offline."),
    ]
    
    for idx, (label, desc) in enumerate(solutions):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        letter = chr(97 + idx)
        r1 = p.add_run(f"({letter}) {label} ")
        set_run_font(r1, size=13, bold=True, color=(31, 78, 121))
        r2 = p.add_run(desc)
        set_run_font(r2, size=13)
    
    add_page_break(doc)
    
    # ====================
    # CHƯƠNG 7: KẾT LUẬN VÀ KIẾN NGHỊ
    # ====================
    add_heading_chapter(doc, "7", "KẾT LUẬN VÀ KIẾN NGHỊ")
    
    add_heading_1(doc, "7.1", "Kết luận")
    
    conclusions = [
        "Dự án đã thành công trong việc xây dựng một pipeline kỹ thuật dữ liệu (ETL) hoàn chỉnh và có chất lượng để phân tích hệ thống giao thông công cộng London. Tất cả mục tiêu đặt ra ban đầu đều được hoàn thành đầy đủ.",
        "Về mặt kỹ thuật, pipeline đã giải quyết thành công bài toán tích hợp dữ liệu dị biệt từ 3 nguồn khác nhau (KML, CSV, NaPTAN) thông qua các thuật toán chuẩn hóa Regex và Nearest Neighbor Matching. Tỷ lệ ghép dữ liệu thành công cao và 298 nhà ga đủ tiêu chuẩn đã được đưa vào phân tích.",
        "Về kết quả phân tích, thuật toán KMeans đã phân loại 298 nhà ga thành 6 cụm có ý nghĩa rõ ràng, phản ánh đúng thực trạng phân cấp của hệ thống giao thông London. Phân tích tác động COVID-19 cung cấp bằng chứng định lượng rõ ràng: lưu lượng giảm 63.21% năm 2020 và đang phục hồi 16.82% năm 2021.",
        "Về sản phẩm cuối, bản đồ tương tác Folium cung cấp giao diện trực quan, thân thiện người dùng với đầy đủ tính năng lọc, tìm kiếm và phân tích. Dữ liệu được lưu trữ đa định dạng và có thể tái sử dụng cho nhiều mục đích phân tích khác nhau.",
        "Dự án đã minh chứng rằng ngay cả với tài nguyên hạn chế (máy tính cá nhân, thư viện Python miễn phí), sinh viên vẫn có thể xây dựng được hệ thống phân tích dữ liệu chuyên nghiệp, có giá trị thực tiễn trên bộ dữ liệu giao thông quy mô thực tế.",
    ]
    for text in conclusions:
        add_body_text(doc, text)
    
    add_heading_1(doc, "7.2", "Hướng phát triển tương lai")
    
    futures = [
        ("Tích hợp dữ liệu thời gian thực:", "Kết nối với TfL Unified API để cập nhật dữ liệu lưu lượng theo giờ/ngày, xây dựng dashboard real-time hiển thị trạng thái mạng lưới tức thì."),
        ("Mô hình dự báo:", "Ứng dụng các mô hình chuỗi thời gian (ARIMA, Prophet, LSTM) để dự báo lưu lượng hành khách cho 2–3 năm tiếp theo, hỗ trợ lập kế hoạch đầu tư cơ sở hạ tầng."),
        ("Mở rộng phân tích địa lý:", "Tích hợp dữ liệu dân số, thu nhập và điểm tắc nghẽn giao thông theo borough để phân tích mối quan hệ giữa nhân khẩu học và nhu cầu giao thông công cộng."),
        ("Phân tích mạng lưới (Graph Analysis):", "Mô hình hóa hệ thống nhà ga như đồ thị (graph) để tìm các nút quan trọng (critical nodes), phân tích tính kết nối và khả năng chịu đựng khi một số nút bị vô hiệu hóa."),
        ("Ứng dụng đề xuất tuyến đường:", "Xây dựng thuật toán đề xuất tuyến đường tối ưu cho hành khách dựa trên lưu lượng dự kiến, tránh các ga quá tải."),
        ("Mở rộng sang các thành phố khác:", "Áp dụng framework ETL đã xây dựng để phân tích dữ liệu giao thông của các thành phố lớn khác như Paris (RATP), Tokyo (Tokyo Metro), New York (MTA)."),
    ]
    
    for label, desc in futures:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        r1 = p.add_run(f"• {label} ")
        set_run_font(r1, size=13, bold=True)
        r2 = p.add_run(desc)
        set_run_font(r2, size=13)
    
    add_page_break(doc)
    
    # ====================
    # TÀI LIỆU THAM KHẢO
    # ====================
    add_paragraph(doc, "TÀI LIỆU THAM KHẢO", align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=14, bold=True, space_after=12)
    
    references = [
        "[1] Transport for London (TfL). (2022). Station Usage Data 2017–2021. Open Data Hub. Truy cập từ https://tfl.gov.uk/corporate/publications-and-reports/",
        "[2] Department for Transport, UK. (2023). NaPTAN — National Public Transport Access Node. Truy cập từ https://www.data.gov.uk/dataset/ff93ffc1-6d58-4a6e-b537-4af99f774e64/national-public-transport-access-nodes-naptan",
        "[3] McKinney, W. (2017). Python for Data Analysis: Data Wrangling with Pandas, NumPy, and IPython (2nd ed.). O'Reilly Media.",
        "[4] Pedregosa, F., Varoquaux, G., Gramfort, A. et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825–2830.",
        "[5] Aggarwal, C. C., & Reddy, C. K. (Eds.). (2014). Data Clustering: Algorithms and Applications. CRC Press.",
        "[6] Python Software Foundation. (2023). Python 3.11 Documentation. Truy cập từ https://docs.python.org/3/",
        "[7] Pandas Development Team. (2023). pandas: Powerful Python Data Analysis Library. Truy cập từ https://pandas.pydata.org/",
        "[8] Scikit-learn Development Team. (2023). sklearn.cluster.KMeans. Truy cập từ https://scikit-learn.org/stable/modules/generated/sklearn.cluster.KMeans.html",
        "[9] Folium Development Team. (2023). Folium Documentation. Truy cập từ https://python-visualization.github.io/folium/",
        "[10] SQLite Consortium. (2023). SQLite Documentation. Truy cập từ https://www.sqlite.org/docs.html",
        "[11] Open Geospatial Consortium. (2008). OGC KML Standard Version 2.2. Truy cập từ https://www.ogc.org/standards/kml/",
        "[12] MacQueen, J. (1967). Some methods for classification and analysis of multivariate observations. Proceedings of the 5th Berkeley Symposium on Mathematical Statistics and Probability, 1, 281–297.",
        "[13] Han, J., Kamber, M., & Pei, J. (2011). Data Mining: Concepts and Techniques (3rd ed.). Morgan Kaufmann.",
        "[14] Leaflet.js Contributors. (2023). Leaflet — a JavaScript library for interactive maps. Truy cập từ https://leafletjs.com/",
        "[15] Bộ Giáo dục và Đào tạo Việt Nam. (2020). Hướng dẫn quy định về trình bày luận văn, đồ án tốt nghiệp. Thông tư số 04/2016/TT-BGDĐT.",
    ]
    
    for ref in references:
        p = add_paragraph(doc, ref, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          size=12, space_after=4, line_spacing=1.5)
    
    add_page_break(doc)
    
        # ====================
    # PHỤ LỤC
    # ====================
    add_paragraph(doc, "PHỤ LỤC", align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=14, bold=True, space_after=12)
    
    add_heading_1(doc, "A.", "Cấu trúc thư mục dự án")
    
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.3
    structure_text = """DE/
├── final_project.py          # Pipeline ETL chính (1,805 dòng)
├── serve_outputs.py          # Web server và tunnel (323 dòng)
├── stations .kml             # Dữ liệu tọa độ nhà ga (~102KB)
├── TfL_stations.csv          # Dữ liệu lưu lượng TfL (~79KB)
├── Stops.csv                 # Dữ liệu NaPTAN (~101MB)
├── FINAL_MAP.html            # Bản đồ tương tác (copy từ outputs)
└── outputs/
    ├── london_tfl_cleaned.csv        # Dữ liệu đã làm sạch
    ├── london_tfl_results.xlsx       # Kết quả Excel (2 sheets)
    ├── london_tfl.db                 # Cơ sở dữ liệu SQLite
    ├── london_tfl_map.html           # Bản đồ tương tác Folium
    └── unmatched_stations_log.txt    # Log các ga không khớp"""
    
    run = p.add_run(structure_text)
    set_run_font(run, size=11)
    
    add_heading_1(doc, "B.", "Lệnh chạy dự án")
    add_body_text(doc, "Để chạy pipeline đầy đủ và tạo tất cả các output:")
    
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run("# Bước 1: Cài đặt các thư viện cần thiết\npip install pandas numpy scikit-learn folium openpyxl reportlab\n\n# Bước 2: Chạy pipeline chính\npython final_project.py\n\n# Bước 3: Khởi động web server để xem bản đồ\npython serve_outputs.py")
    set_run_font(run, size=11)
    
    add_heading_1(doc, "C.", "Tài liệu kiểm toán kỹ thuật và sửa lỗi đồng bộ")
    add_appendix_c_content(doc)
    
    add_heading_1(doc, "D.", "Tổng hợp tài liệu thuật toán và cấu trúc hệ thống")
    add_appendix_d_from_markdown(doc, "tong_hop_ky_thuat.md")
    
    add_heading_1(doc, "E.", "Log các ga không khớp")
    add_body_text(doc, "File outputs/unmatched_stations_log.txt ghi lại danh sách các nhà ga trong file KML không ghép được với dữ liệu hành khách TfL hoặc bị loại bỏ do thiếu thông tin. Các lý do phổ biến: tên ga quá khác biệt giữa các nguồn sau khi chuẩn hóa, nhà ga chưa có dữ liệu lưu lượng (ga mới hoặc đã đóng cửa), và thiếu tọa độ GPS hợp lệ.")
    
    # Lưu file
    output_path = Path("c:/Users/My PC/Downloads/DE/BaoCao_TfL_London_Nhom14_Final.docx")
    doc.save(str(output_path))
    print(f"✅ Đã tạo FILE 1: {output_path}")
    return output_path


# ============================================================
# FILE 2: TÀI LIỆU PHÂN TÍCH CODE
# ============================================================

def create_code_analysis():
    doc = Document()
    set_page_margins(doc, top=2, bottom=2, left=3, right=2)
    
    # TRANG BÌA
    for text, sz, bold in [
        ("TRƯỜNG ĐẠI HỌC CÔNG NGHIỆP TP.HCM", 13, True),
        ("KHOA CÔNG NGHỆ THÔNG TIN", 13, True),
    ]:
        add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, size=sz, bold=bold, space_after=4)
    
    add_horizontal_line(doc)
    
    for _ in range(4):
        add_paragraph(doc, "")
    
    add_paragraph(doc, "TÀI LIỆU PHÂN TÍCH CODE", align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=18, bold=True, color=(31, 78, 121), space_after=8)
    add_paragraph(doc, "Dự án: London Transport Analysis (TfL)", align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=14, bold=True, space_after=4)
    add_paragraph(doc, "Phân tích chi tiết codebase — Nhóm 14", align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=13, italic=True, space_after=4)
    
    for _ in range(3):
        add_paragraph(doc, "")
    
    add_paragraph(doc, "File phân tích: final_project.py (1,805 dòng) & serve_outputs.py (323 dòng)",
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=12, italic=True)
    
    for _ in range(4):
        add_paragraph(doc, "")
    add_paragraph(doc, "TP. HỒ CHÍ MINH, THÁNG 05 NĂM 2026",
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True)
    
    add_page_break(doc)
    
    # ====================
    # PHẦN 1: GIỚI THIỆU
    # ====================
    add_heading_chapter(doc, "1", "GIỚI THIỆU")
    
    add_heading_1(doc, "1.1", "Mục đích tài liệu")
    add_body_text(doc, "Tài liệu này cung cấp phân tích kỹ thuật chuyên sâu về codebase của dự án TfL London Analysis, bao gồm hai file Python chính: final_project.py (pipeline ETL và phân tích chính) và serve_outputs.py (web server và tunneling). Mục tiêu là cung cấp hiểu biết đầy đủ về kiến trúc code, các quyết định thiết kế, và hướng dẫn vận hành chi tiết cho người đọc có kiến thức Python trung cấp trở lên.")
    
    add_heading_1(doc, "1.2", "Tổng quan codebase")
    
    overview_data = [
        ("final_project.py", "1,805", "91,409 bytes (~89KB)", "Pipeline ETL, phân cụm, trực quan hóa"),
        ("serve_outputs.py", "323", "10,886 bytes (~11KB)", "Web server, ngrok/localtunnel integration"),
        ("generate_pdf_report.py", "~400", "~41KB", "Xuất báo cáo PDF"),
        ("create_presentation.py", "~600", "~67KB", "Tạo slide trình bày PowerPoint"),
    ]
    add_table_with_caption(doc, "Bảng 1.1: Tổng quan các file code của dự án",
                           ["File", "Số dòng", "Kích thước", "Chức năng chính"],
                           overview_data, col_widths=[4.5, 2, 3, 6.5])
    
    add_page_break(doc)
    
    # ====================
    # PHẦN 2: PHÂN TÍCH final_project.py
    # ====================
    add_heading_chapter(doc, "2", "PHÂN TÍCH CHI TIẾT: final_project.py")
    
    add_heading_1(doc, "2.1", "Cấu trúc tổng thể và kiến trúc code")
    add_body_text(doc, "File final_project.py được tổ chức theo cấu trúc module rõ ràng với 7 section chính, đánh số từ 0 đến 6 và ngăn cách bằng các comment block dạng '='*72. Đây là một pattern phổ biến trong code khoa học dữ liệu, giúp người đọc nhanh chóng định vị vùng code cần quan tâm.")
    
    sections = [
        ("Section 0 — CẤU HÌNH CHUNG (dòng 53–112):", "Định nghĩa các hằng số, danh sách cấu hình và dataclass. Thiết kế tập trung này đảm bảo tất cả tham số có thể tùy chỉnh ở một nơi duy nhất mà không cần sửa logic."),
        ("Section 1 — HÀM HỖ TRỢ (dòng 126–210):", "Các utility function dùng chung: get_project_folder(), find_file(), ensure_package(), normalize_station_name(), first_non_empty(), print_title()."),
        ("Section 2 — EXTRACT (dòng 216–380):", "Ba hàm đọc dữ liệu: load_kml_stations(), load_tfl_csv(), load_stops_csv(). Mỗi hàm xử lý một định dạng file khác nhau và trả về DataFrame chuẩn hóa."),
        ("Section 3 — TRANSFORM (dòng 386–541):", "Các hàm biến đổi dữ liệu: nearest_stop_match(), merge_sources(), clean_and_engineer(), add_trend_analysis()."),
        ("Section 4 — CLUSTERING & ANALYSIS (dòng 547–724):", "Các hàm phân tích: run_kmeans_clustering(), create_cluster_summary(), print_analysis_report(), build_summary_text()."),
        ("Section 5 — LOAD (dòng 729–788):", "Các hàm lưu trữ: save_outputs(), log_unmatched_stations(), make_sqlite_safe_dataframe()."),
        ("Section 6 — VISUALIZATION (dòng 793–1734):", "Hàm create_folium_map() — phần lớn nhất (khoảng 940 dòng) chứa toàn bộ code HTML/CSS/JavaScript cho bản đồ tương tác."),
        ("Section 7 — MAIN (dòng 1740–1805):", "Hàm run_pipeline() điều phối toàn bộ pipeline và khối if __name__ == '__main__'."),
    ]
    
    for label, desc in sections:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        r1 = p.add_run(f"▶ {label} ")
        set_run_font(r1, size=13, bold=True)
        r2 = p.add_run(desc)
        set_run_font(r2, size=13)
    
    add_heading_1(doc, "2.2", "Section 0 — Cấu hình và Constants")
    
    add_heading_2(doc, "2.2.1", "Hằng số và danh sách cấu hình")
    add_body_text(doc, "Tất cả tham số có khả năng thay đổi được tập trung ở đầu file, theo nguyên tắc Single Source of Truth:")
    
    constants_code = """KML_CANDIDATES = ["stations .kml", "stations.kml"]  # Tên file KML có thể có
TFL_CANDIDATES = ["TfL_stations.csv", "stations.csv"]  # Tên file TfL CSV
STOPS_CANDIDATES = ["Stops.csv"]  # File NaPTAN

YEARS = [2017, 2018, 2019, 2020, 2021]  # Năm phân tích
N_CLUSTERS = 6  # Số cụm KMeans

LINE_SUFFIXES = [  # Danh sách hậu tố cần loại bỏ khi chuẩn hóa tên ga
    " underground station", " overground station", " rail station",
    " dlr station", " elizabeth line station", " tram stop", " station",
    " lu", " lo", " nr", " dlr", " el", " tfl", " (dis)", ...
]

STOPTYPE_PRIORITY = {  # Thứ tự ưu tiên chọn record từ NaPTAN
    "MET": 1, "RLY": 2, "RSE": 3, "TMU": 4, "DLR": 5, "PLT": 6, "RPL": 7
}

CLUSTER_NAMES = [  # Tên hiển thị cho 6 cụm theo thứ tự lưu lượng giảm dần
    "Siêu trung tâm", "Ga lớn", "Ga trung bình",
    "Ga nhỏ", "Ga ít khách", "Ga rất ít khách"
]"""
    
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(constants_code)
    set_run_font(run, "Courier New", size=10)
    
    add_heading_2(doc, "2.2.2", "Dataclass MergeReport")
    add_body_text(doc, "MergeReport là Python dataclass được sử dụng để đóng gói các thống kê về quá trình ghép dữ liệu (số dòng từng nguồn, số khớp thành công). Sử dụng dataclass thay vì dict thông thường tăng tính rõ ràng (type hints) và tránh lỗi do gõ sai key. Pattern này là best practice trong Python 3.7+ cho việc truyền dữ liệu có cấu trúc giữa các hàm.")
    
    add_heading_1(doc, "2.3", "Section 1 — Hàm hỗ trợ")
    
    add_heading_2(doc, "2.3.1", "normalize_station_name() — Cốt lõi của Data Cleaning")
    add_body_text(doc, "Đây là hàm quan trọng nhất và tinh tế nhất trong codebase. Hàm thực hiện quy trình chuẩn hóa 5 bước với vòng lặp while để xử lý nhiều hậu tố lồng nhau (ví dụ: 'bank underground station lu' → 'bank' sau nhiều vòng lặp).")
    
    normalize_code = """def normalize_station_name(name: object) -> str:
    if pd.isna(name): return ""
    
    text = str(name).strip().lower()
    text = text.replace("&", " and").replace("st.", "st")
    text = re.sub(r"\\s+", " ", text)

    # Vòng lặp xử lý nhiều hậu tố lồng nhau
    changed = True
    while changed:
        changed = False
        for suffix in LINE_SUFFIXES:
            if text.endswith(suffix):
                text = text[: -len(suffix)].strip()
                changed = True
                break  # Restart loop khi tìm thấy suffix

    text = re.sub(r"[^\\w\\s']", " ", text)  # Xóa ký tự đặc biệt
    text = re.sub(r"\\s+", " ", text).strip()

    # Manual aliases cho trường hợp đặc biệt
    if text in ["bank", "monument"]:
        return "bank and monument"
    ...
    return text"""
    
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(normalize_code)
    set_run_font(run, "Courier New", size=10)
    
    add_body_text(doc, "Điểm đáng chú ý về thiết kế: Vòng lặp while thay vì for đơn lẻ đảm bảo tất cả hậu tố được loại bỏ khi chúng xuất hiện kết hợp. Danh sách LINE_SUFFIXES được thiết kế theo thứ tự từ cụ thể nhất đến tổng quát nhất (ví dụ: ' underground station' trước ' station') để tránh loại bỏ thiếu. Các alias thủ công xử lý những trường hợp mà thuật toán chung không thể giải quyết.")
    
    add_heading_2(doc, "2.3.2", "ensure_package() — Auto-Install Dependencies")
    add_body_text(doc, "Hàm này tự động cài đặt thư viện Python còn thiếu qua pip trong runtime. Đây là pattern hữu ích khi chạy trên Google Colab — nơi người dùng không thể chạy pip trước khi chạy code. Tuy nhiên, pattern này không được khuyến nghị trong production do có thể gây ra vấn đề về version compatibility và security. Best practice trong production là sử dụng requirements.txt hoặc pyproject.toml.")
    
    add_heading_1(doc, "2.4", "Section 2 — Extract")
    
    add_heading_2(doc, "2.4.1", "load_kml_stations() — Đọc file KML")
    add_body_text(doc, "Thách thức kỹ thuật chính: File KML thực tế có XML namespace (thẻ có dạng {http://www.opengis.net/kml/2.2}name), làm cho XPath trực tiếp thất bại. Giải pháp: Strip toàn bộ namespace trước khi xử lý bằng vòng lặp elem.tag.split('}', 1)[1].")
    
    kml_code = """def load_kml_stations(kml_path: Path) -> pd.DataFrame:
    tree = ET.parse(kml_path)
    root = tree.getroot()

    # Strip namespaces (giải quyết XML namespace problem)
    for elem in root.iter():
        if elem.tag.startswith('{'):
            elem.tag = elem.tag.split('}', 1)[1]  # Chỉ giữ tên tag, bỏ namespace

    rows = []
    for placemark in root.findall(".//Placemark"):  # XPath tìm tất cả Placemark
        name_el = placemark.find("name")
        coord_el = placemark.find(".//coordinates")

        # Kiểm tra đầy đủ trước khi parse
        if name_el is None or coord_el is None: continue
        
        parts = coord_el.text.strip().split(",")  # Format: "lon,lat,alt"
        lon, lat = float(parts[0]), float(parts[1])

        rows.append({
            "station": station,
            "station_key": normalize_station_name(station),  # Key dùng để merge
            "lat": lat, "lon": lon
        })

    return pd.DataFrame(rows).drop_duplicates(subset=["station_key"])\
                             .reset_index(drop=True)"""
    
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(kml_code)
    set_run_font(run, "Courier New", size=10)
    
    add_heading_2(doc, "2.4.2", "load_tfl_csv() — Đọc file CSV đặc biệt")
    add_body_text(doc, "File TfL_stations.csv có bug định dạng: mỗi dòng dữ liệu bị bọc trong dấu ngoặc kép như một string đơn. Hàm xử lý bằng cách parse hai lần: lần 1 parse CSV bình thường → nếu chỉ ra 1 phần tử → parse lại chuỗi đó như CSV. Logic groupby/agg sau đó gộp các nhà ga có cùng station_key (như Liverpool Street LU và NR) thành một hàng duy nhất bằng cách cộng tổng lưu lượng hành khách.")
    
    add_heading_2(doc, "2.4.3", "load_stops_csv() — Đọc file NaPTAN lớn")
    add_body_text(doc, "File 101MB được đọc với low_memory=False để tránh lỗi mixed-type inference. Trường Borough được enrich từ nhiều cột dự phòng (Town > ParentLocalityName > LocalityName) theo thứ tự ưu tiên bằng hàm first_non_empty(). Cột stop_priority được tính theo bảng STOPTYPE_PRIORITY để sau này dùng trong Nearest Neighbor selection.")
    
    add_heading_1(doc, "2.5", "Section 3 — Transform")
    
    add_heading_2(doc, "2.5.1", "nearest_stop_match() — Spatial Join với Nearest Neighbor")
    add_body_text(doc, "Hàm này giải quyết vấn đề một tên ga có thể match nhiều record trong NaPTAN. Thuật toán hoạt động theo 3 bước: (1) LEFT JOIN theo station_key để lấy tất cả candidates; (2) Tính khoảng cách Euclidean giữa tọa độ KML và từng NaPTAN record; (3) Sắp xếp theo [stop_priority ASC, stop_distance ASC] rồi lấy record đầu tiên (best match).")
    
    spatial_code = """def nearest_stop_match(merged_df, stops_df):
    candidates = pd.merge(
        merged_df[["station", "station_key", "lat", "lon"]],
        stops_df[["station_key", "CommonName", "borough", 
                  "stop_lat", "stop_lon", "StopType", "stop_priority"]],
        on="station_key", how="left"  # LEFT JOIN giữ tất cả nhà ga
    )
    
    # Khoảng cách Euclidean (không cần Haversine do scale nhỏ trong London)
    candidates["stop_distance"] = np.sqrt(
        (candidates["lat"] - candidates["stop_lat"]) ** 2
        + (candidates["lon"] - candidates["stop_lon"]) ** 2
    )
    
    # Sắp xếp: ưu tiên MET trước, sau đó theo khoảng cách
    candidates = candidates.sort_values(["station", "stop_priority", "stop_distance"])
    best = candidates.drop_duplicates(subset=["station"], keep="first")  # Best match mỗi ga
    
    result = pd.merge(merged_df, 
                      best[["station", "borough", "CommonName", "StopType", "stop_distance"]],
                      on="station", how="left")
    return result.rename(columns={"CommonName": "stop_name", "StopType": "stop_type"})"""
    
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(spatial_code)
    set_run_font(run, "Courier New", size=10)
    
    add_body_text(doc, "Lý do dùng Euclidean thay vì Haversine: Vì phạm vi địa lý nhỏ (toàn bộ trong London, bán kính ~25km), sự sai lệch giữa hai công thức là không đáng kể (~0.1%) và Euclidean nhanh hơn nhiều khi tính trên hàng triệu cặp điểm.")
    
    add_heading_2(doc, "2.5.2", "clean_and_engineer() — Feature Engineering")
    add_body_text(doc, "Hàm này tạo ra 6 đặc trưng phân tích mới và thực hiện lọc cuối cùng để chỉ giữ lại các nhà ga đủ điều kiện (có tọa độ hợp lệ, có dữ liệu hành khách 2021 và thông tin tuyến).")
    
    features_code = """def clean_and_engineer(df):
    # Chuyển đổi cột "En/Ex YYYY" sang passengers_YYYY (số nguyên)
    for year in YEARS:
        raw_col = f"En/Ex {year}"
        new_col = f"passengers_{year}"
        df[new_col] = pd.to_numeric(
            df[raw_col].astype(str).str.replace(",", "").str.strip(),
            errors="coerce"  # Không hợp lệ → NaN (không crash)
        )
    
    # Feature: Số tuyến phục vụ
    df["num_lines"] = df["LINES"].fillna("").str.count(",") + 1
    
    # Feature: COVID Impact (BUG FIX #3: Tránh ZeroDivisionError)
    df["covid_impact_pct"] = np.where(
        (df["passengers_2019"] == 0) | df["passengers_2019"].isna(),
        np.nan,  # Nếu 2019 = 0 hoặc NaN, không tính
        (df["passengers_2020"] - df["passengers_2019"]) / df["passengers_2019"] * 100
    )
    
    # Lọc cuối: Chỉ giữ ga có đủ thông tin
    df = df.dropna(subset=["lat", "lon", "passengers_2021", "num_lines"])
    return df"""
    
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(features_code)
    set_run_font(run, "Courier New", size=10)
    
    add_heading_2(doc, "2.5.3", "add_trend_analysis() — LinearRegression cho Trend")
    add_body_text(doc, "Hàm tính xu hướng 5 năm của từng nhà ga bằng LinearRegression từ sklearn. BUG FIX #4 quan trọng: Thay vì bỏ qua toàn bộ hàng nếu có một năm NaN, code lọc ra các năm có dữ liệu hợp lệ (không NaN, không 0) và fit model chỉ trên các điểm hợp lệ đó. Điều này cho phép tính xu hướng ngay cả khi thiếu dữ liệu 1–2 năm, miễn là còn ít nhất 2 điểm.")
    
    add_heading_1(doc, "2.6", "Section 4 — Clustering và Analysis")
    
    add_heading_2(doc, "2.6.1", "run_kmeans_clustering() — Phân cụm")
    
    kmeans_code = """def run_kmeans_clustering(df, n_clusters=N_CLUSTERS):
    features = ["passengers_2021", "num_lines", "lat", "lon"]
    matrix = df[features].to_numpy(dtype=float)
    
    # Chuẩn hóa về phân phối chuẩn (z-score normalization)
    matrix_scaled = StandardScaler().fit_transform(matrix)
    
    # BUG FIX #5: Giảm số cụm nếu dữ liệu quá ít
    actual_clusters = min(n_clusters, len(df))
    
    model = KMeans(
        n_clusters=actual_clusters,
        random_state=42,  # Đảm bảo reproducibility
        n_init=10         # Chạy 10 lần với init khác nhau, lấy kết quả tốt nhất
    )
    df["cluster_id"] = model.fit_predict(matrix_scaled)
    
    # Đặt tên cụm theo thứ tự lưu lượng trung bình giảm dần
    cluster_means = df.groupby("cluster_id")["passengers_2021"].mean()\
                      .sort_values(ascending=False)
    name_map = {cid: CLUSTER_NAMES[rank] 
                for rank, cid in enumerate(cluster_means.index)}
    df["cluster_name"] = df["cluster_id"].map(name_map)
    return df"""
    
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(kmeans_code)
    set_run_font(run, "Courier New", size=10)
    
    add_body_text(doc, "Lý do chọn StandardScaler: Đặc trưng passengers_2021 có giá trị hàng triệu trong khi lat/lon chỉ có giá trị từ 51–52 (độ vĩ) và -0.5 đến 0.3 (độ kinh). Nếu không chuẩn hóa, KMeans sẽ bị dominated bởi passengers và gần như bỏ qua yếu tố vị trí địa lý, dẫn đến phân cụm sai về mặt địa lý.")
    
    add_heading_1(doc, "2.7", "Section 5 — Load")
    
    add_heading_2(doc, "2.7.1", "make_sqlite_safe_dataframe() — Xử lý SQLite Constraints")
    add_body_text(doc, "Hàm này giải quyết vấn đề thú vị: SQLite không phân biệt chữ hoa/thường trong tên cột, nên 'Station' và 'station' bị coi là trùng nhau khi tạo bảng. Hàm chuyển tất cả tên cột sang lowercase, thay ký tự đặc biệt bằng '_' và tự động đặt lại tên nếu trùng (thêm hậu tố _2, _3...).")
    
    add_heading_2(doc, "2.7.2", "save_outputs() — Multi-format Export")
    add_body_text(doc, "Hàm xuất dữ liệu ra 4 định dạng đồng thời: CSV (to_csv), Excel (ExcelWriter với 2 sheets), SQLite (to_sql với bảng fact_stations và dim_clusters) và Text report (write_text). Việc kiểm tra và tự động cài đặt openpyxl trong runtime (ensure_package('openpyxl')) cho phép Excel export chạy ngay cả khi người dùng quên cài thư viện này.")
    
    add_heading_1(doc, "2.8", "Section 6 — Visualization Module")
    
    add_heading_2(doc, "2.8.1", "create_folium_map() — Embedded Web App")
    add_body_text(doc, "Đây là phần phức tạp nhất của codebase (940 dòng). Thay vì dùng API Folium Python để tạo bản đồ đơn giản, code nhúng toàn bộ một ứng dụng web JavaScript/HTML tùy chỉnh vào một template string và inject dữ liệu bằng cách replace placeholder (__STATION_DATA__, __CENTER_LAT__, v.v.).")
    
    add_body_text(doc, "Kiến trúc JavaScript của bản đồ bao gồm các thành phần chính:")
    
    js_components = [
        ("Data Layer:", "Dữ liệu nhà ga được JSON.stringify() từ Python và nhúng trực tiếp vào HTML dưới dạng JavaScript constant. Cách này không cần AJAX request và bản đồ hoạt động offline ngoại trừ phần tile."),
        ("MarkerCluster Layer:", "Sử dụng Leaflet.MarkerCluster với custom iconCreateFunction hiển thị tổng lượt hành khách trong từng cluster. Tự động giải phóng (spiderfy) khi zoom vào mức 13+."),
        ("Filter System:", "applyFilters() là hàm trung tâm được gọi mỗi khi bất kỳ bộ lọc nào thay đổi. Hàm duyệt toàn bộ markerList và quyết định show/hide từng marker dựa trên 3 điều kiện: cluster_category, borough và search_term."),
        ("Analytics Engine:", "updateAnalytics() tính toán thống kê realtime chỉ cho các marker đang hiển thị trong viewport hiện tại (getBounds().contains()), cập nhật top-stations list và biểu đồ phân bố cụm."),
        ("Heatmap Layer:", "buildHeatLayer() tạo Leaflet.heat với gradient từ xanh → vàng → đỏ, sử dụng căn bậc hai của intensity để làm nổi bật các ga nhỏ (tránh bị che khuất bởi các ga siêu lớn)."),
        ("Sparkline SVG:", "renderSparkline() tạo biểu đồ đường SVG nhỏ (140×36px) hiển thị xu hướng 5 năm của từng nhà ga trong popup."),
    ]
    
    for idx, (label, desc) in enumerate(js_components):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        letter = chr(97 + idx)
        r1 = p.add_run(f"({letter}) {label} ")
        set_run_font(r1, size=13, bold=True)
        r2 = p.add_run(desc)
        set_run_font(r2, size=13)
    
    add_heading_2(doc, "2.8.2", "Error Handling và các BUG FIX")
    
    bugfixes = [
        ("BUG FIX #1 — Empty merge check (dòng 403):", "Thay vì kiểm tra len(candidates) == 0, code kiểm tra candidates['CommonName'].isna().all() — chính xác hơn với LEFT JOIN vì kết quả luôn có dòng nhưng có thể toàn NaN."),
        ("BUG FIX #2 — CSV double-parsing logic (dòng 293):", "Cải thiện logic xử lý CSV wrapping: chỉ parse lại khi len(first_pass) == 1 VÀ có dấu phẩy trong chuỗi. Điều kiện kép tránh parse sai các ô text đơn không chứa CSV data."),
        ("BUG FIX #3 — ZeroDivisionError trong COVID impact (dòng 477):", "Dùng np.where() với điều kiện kiểm tra (passengers_2019 == 0) | passengers_2019.isna() thay vì chia trực tiếp. Vectorized và an toàn hơn try/except."),
        ("BUG FIX #4 — NaN trong Trend Analysis (dòng 511):", "Lọc valid_mask = ~np.isnan(passengers) & (passengers > 0) để chỉ fit LinearRegression trên các năm có dữ liệu thực. Yêu cầu tối thiểu 2 điểm (min 2 năm) để fit đường thẳng."),
        ("BUG FIX #5 — Insufficient data for KMeans (dòng 560):", "actual_clusters = min(n_clusters, len(df)) đảm bảo không yêu cầu nhiều cụm hơn số điểm dữ liệu. Kèm cảnh báo rõ ràng ra console."),
    ]
    
    for idx, (label, desc) in enumerate(bugfixes):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        letter = chr(97 + idx)
        r1 = p.add_run(f"({letter}) {label} ")
        set_run_font(r1, size=13, bold=True, color=(192, 80, 77))
        r2 = p.add_run(desc)
        set_run_font(r2, size=13)
    
    add_page_break(doc)
    
    # ====================
    # PHẦN 3: PHÂN TÍCH serve_outputs.py
    # ====================
    add_heading_chapter(doc, "3", "PHÂN TÍCH CHI TIẾT: serve_outputs.py")
    
    add_heading_1(doc, "3.1", "Kiến trúc Web Server")
    add_body_text(doc, "serve_outputs.py là module độc lập đóng vai trò là web server và network tunnel manager, cho phép chia sẻ bản đồ HTML ra ngoài internet. File sử dụng Python standard library http.server kết hợp với các công cụ tunnel bên ngoài (ngrok, localtunnel).")
    
    server_flow = [
        ("Bước 1 — Dependency Check:", "Kiểm tra và cài đặt tự động các Python package cần thiết (pandas, numpy, sklearn, requests, pyngrok) bằng hàm ensure_python_packages()."),
        ("Bước 2 — Pipeline Execution:", "Gọi run_pipeline() để chạy lại final_project.py và tạo lại file HTML mới nhất (nếu dữ liệu đã thay đổi)."),
        ("Bước 3 — Port Discovery:", "Thử lần lượt từ PORT=8000, tăng dần nếu bị chiếm dụng (tối đa 20 lần thử) bằng cách bắt OSError."),
        ("Bước 4 — HTTP Server:", "Khởi động ThreadingHTTPServer phục vụ static files từ thư mục outputs/. Chạy trong daemon thread để không block main thread."),
        ("Bước 5 — Tunnel Creation:", "Thử tạo tunnel theo thứ tự: ngrok (nếu có NGROK_AUTH_TOKEN) → localtunnel (qua npx). Cả hai đều tạo public HTTPS URL."),
        ("Bước 6 — Keep-alive Loop:", "Main thread chạy while True với sleep(1) để giữ server hoạt động. Ctrl+C kích hoạt cleanup (tắt server, ngắt tunnel)."),
    ]
    
    for idx, (label, desc) in enumerate(server_flow):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        letter = chr(97 + idx)
        r1 = p.add_run(f"({letter}) {label} ")
        set_run_font(r1, size=13, bold=True)
        r2 = p.add_run(desc)
        set_run_font(r2, size=13)
    
    add_heading_1(doc, "3.2", "Localtunnel/ngrok Integration")
    
    add_heading_2(doc, "3.2.1", "create_public_tunnel() — ngrok")
    add_body_text(doc, "Hàm này tạo ngrok tunnel thông qua thư viện pyngrok. ngrok yêu cầu auth token (lấy miễn phí tại ngrok.com) được lưu trong biến môi trường NGROK_AUTH_TOKEN. Ưu điểm: URL ổn định, hỗ trợ HTTPS, có dashboard web. Nhược điểm: Cần đăng ký tài khoản, giới hạn 1 tunnel với tài khoản free.")
    
    add_heading_2(doc, "3.2.2", "create_localtunnel() — Fallback option")
    add_body_text(doc, "Localtunnel (lt) là giải pháp thay thế không cần đăng ký, sử dụng qua npx (Node.js). Hàm đọc output của process localtunnel theo pattern regex 'your url is: (https?://...)' với timeout 20 giây. Cơ chế fallback: thử subdomain cố định trước (dựa trên hash đường dẫn) → nếu thất bại, thử lại với subdomain ngẫu nhiên.")
    
    tunnel_code = """def create_localtunnel(port: int):
    # Tạo subdomain cố định từ hash đường dẫn (tránh URL thay đổi mỗi lần)
    default_subdomain = f"london-tfl-{hashlib.sha1(str(OUTPUT_DIR).encode()).hexdigest()[:8]}"
    
    cmd = ["npx", "--yes", "localtunnel", "--port", str(port),
           "--print-requests", "--subdomain", default_subdomain]
    
    # Windows cần wrapper cmd /c
    if sys.platform.startswith("win"):
        cmd = ["cmd", "/c"] + cmd
    
    proc = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT)
    
    # Đọc output theo dòng với timeout 20 giây
    url_pattern = re.compile(r"your url is: (https?://[\\w\\.-]+)", re.IGNORECASE)
    while time.time() - start_time < 20:
        line = proc.stdout.readline()
        if match := url_pattern.search(line):
            url = match.group(1)
            break
    
    return proc  # Giữ process để tunnel không bị đóng"""
    
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(tunnel_code)
    set_run_font(run, "Courier New", size=10)
    
    add_heading_1(doc, "3.3", "Cách vận hành và Output")
    add_body_text(doc, "Khi chạy thành công, serve_outputs.py in ra 4 URL truy cập:")
    
    urls = [
        ("Local:", "http://127.0.0.1:8000/london_tfl_map.html — Chỉ truy cập từ máy đang chạy"),
        ("LAN:", "http://192.168.x.x:8000/london_tfl_map.html — Truy cập từ thiết bị cùng mạng WiFi"),
        ("PDF Local:", "http://127.0.0.1:8000/TfL_Project_Report.pdf — Báo cáo PDF"),
        ("Public:", "https://london-tfl-xxxxxxxx.loca.lt/london_tfl_map.html — Truy cập từ internet"),
    ]
    for label, desc in urls:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        r1 = p.add_run(f"• {label} ")
        set_run_font(r1, size=13, bold=True)
        r2 = p.add_run(desc)
        set_run_font(r2, size=12)
    
    add_page_break(doc)
    
    # ====================
    # PHẦN 4: ĐÁNH GIÁ CODEBASE
    # ====================
    add_heading_chapter(doc, "4", "ĐÁNH GIÁ TỔNG THỂ CODEBASE")
    
    add_heading_1(doc, "4.1", "Điểm mạnh")
    
    strengths = [
        "Code Python idioms tốt: Sử dụng dataclass, type hints (list[str], Path, pd.DataFrame), f-strings và walrus operator (:=) — đặc trưng của Python 3.10+.",
        "Defensive programming: Kiểm tra input, xử lý lỗi với try/except, không để crash không có thông báo.",
        "Single Responsibility: Mỗi hàm có một nhiệm vụ rõ ràng, hàm dài nhất là create_folium_map() nhưng tính phức tạp là không thể tránh khỏi.",
        "Logging tốt: In progress report rõ ràng tại mỗi bước, giúp debug dễ dàng.",
        "UTF-8 handling: Xử lý encoding đặc biệt cho Windows console (sys.stdout.reconfigure).",
        "Reproducibility: random_state=42 trong KMeans đảm bảo kết quả giống nhau mỗi lần chạy.",
    ]
    for idx, s in enumerate(strengths):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        letter = chr(97 + idx)
        run = p.add_run(f"({letter}) {s}")
        set_run_font(run, size=13)
    
    add_heading_1(doc, "4.2", "Điểm cần cải thiện")
    
    weaknesses = [
        "Thiếu unit test: Không có pytest test suite. Khó đảm bảo code không bị regression khi refactor.",
        "Quá trình tạo HTML quá dài: 940 dòng HTML/CSS/JS nhúng trong Python string gây khó đọc. Nên tách ra file template riêng.",
        "N+1 problem trong trend analysis: Vòng lặp for _, row in df.iterrows() với LinearRegression bên trong — chậm với DataFrame lớn. Nên vectorize bằng apply() hoặc numpy broadcasting.",
        "Magic numbers: Một số giá trị như minRadius=6, maxRadius=28, heat radius=28 không được định nghĩa là hằng số có tên, khó hiểu ý nghĩa.",
        "Thiếu logging module: Dùng print() thay vì Python logging module, không có log level (DEBUG/INFO/WARNING/ERROR).",
    ]
    for idx, w in enumerate(weaknesses):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        letter = chr(97 + idx)
        run = p.add_run(f"({letter}) {w}")
        set_run_font(run, size=13)
    
    add_heading_1(doc, "4.3", "Gợi ý nâng cấp Production")
    
    upgrades = [
        ("Containerization:", "Đóng gói bằng Docker để đảm bảo môi trường nhất quán. Tạo Dockerfile với Python image và requirements.txt."),
        ("CI/CD Pipeline:", "Tích hợp GitHub Actions để tự động chạy tests và deploy khi push code mới."),
        ("Database upgrade:", "Chuyển từ SQLite sang PostgreSQL/DuckDB cho production với nhiều người dùng đồng thời."),
        ("Caching:", "Cache kết quả KMeans và merge vào disk để không phải tính lại từ đầu khi dữ liệu không thay đổi."),
        ("API REST:", "Thay vì serve static HTML, xây dựng FastAPI backend cung cấp API endpoint /stations, /clusters, /stats để frontend có thể query động."),
        ("Monitoring:", "Tích hợp Prometheus + Grafana để theo dõi performance và usage metrics của web server."),
    ]
    for idx, (label, desc) in enumerate(upgrades):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        letter = chr(97 + idx)
        r1 = p.add_run(f"({letter}) {label} ")
        set_run_font(r1, size=13, bold=True)
        r2 = p.add_run(desc)
        set_run_font(r2, size=13)
    
    add_page_break(doc)
    
    # ====================
    # PHẦN 5: HƯỚNG DẪN VẬN HÀNH
    # ====================
    add_heading_chapter(doc, "5", "HƯỚNG DẪN VẬN HÀNH DỰ ÁN")
    
    add_heading_1(doc, "5.1", "Yêu cầu môi trường")
    
    env_reqs = [
        ("Hệ điều hành:", "Windows 10/11, macOS 12+, Ubuntu 20.04+ hoặc Google Colab"),
        ("Python:", "3.10 trở lên (yêu cầu type hint union syntax list[str] và match statement)"),
        ("RAM:", "Tối thiểu 4GB (Stops.csv 101MB cần ~2GB RAM khi xử lý)"),
        ("Dung lượng ổ cứng:", "Tối thiểu 500MB (data + outputs)"),
        ("Kết nối internet:", "Cần thiết cho bản đồ (CDN tiles) và localtunnel"),
    ]
    for idx, (label, desc) in enumerate(env_reqs):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        letter = chr(97 + idx)
        r1 = p.add_run(f"({letter}) {label} ")
        set_run_font(r1, size=13, bold=True)
        r2 = p.add_run(desc)
        set_run_font(r2, size=13)
    
    add_heading_1(doc, "5.2", "Hướng dẫn cài đặt và chạy")
    
    install_steps = [
        ("Bước 1 — Cài đặt Python:", "Tải Python 3.10+ từ python.org. Đảm bảo tick 'Add to PATH' khi cài."),
        ("Bước 2 — Cài thư viện:", "Mở terminal trong thư mục dự án và chạy:\n  pip install pandas numpy scikit-learn folium openpyxl reportlab python-pptx"),
        ("Bước 3 — Đặt file dữ liệu:", "Đảm bảo 3 file dữ liệu (stations .kml, TfL_stations.csv, Stops.csv) nằm trong cùng thư mục với final_project.py"),
        ("Bước 4 — Chạy pipeline:", "python final_project.py\n  Pipeline sẽ tự động chạy qua 6 bước và tạo thư mục outputs/"),
        ("Bước 5 — Xem bản đồ:", "Mở file outputs/london_tfl_map.html bằng trình duyệt web, hoặc chạy:\n  python serve_outputs.py\n  để phục vụ qua web server và có public URL"),
    ]
    for idx, (label, desc) in enumerate(install_steps):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        letter = chr(97 + idx)
        r1 = p.add_run(f"({letter}) {label}\n")
        set_run_font(r1, size=13, bold=True)
        r2 = p.add_run(f"   {desc}")
        set_run_font(r2, size=12)
    
    add_heading_1(doc, "5.3", "Output mong đợi và thời gian chạy")
    
    expected_outputs = [
        ("KML: doc duoc X ga...", "Số nhà ga đọc được từ file KML"),
        ("TfL CSV: doc va gop luu luong duoc X ga...", "Số nhà ga unique sau gộp"),
        ("Stops CSV: doc duoc X ban ghi hop le", "Số dòng hợp lệ từ NaPTAN"),
        ("Sau lam sach: giu lai 298/... ga...", "Số ga đủ điều kiện phân tích"),
        ("Cluster 0-5 — [Tên cụm]: X ga...", "Tóm tắt từng cụm KMeans"),
        ("Da luu CSV/Excel/SQLite/Ban do...", "Xác nhận lưu output thành công"),
    ]
    add_table_with_caption(doc, "Bảng 5.1: Output mong đợi khi chạy final_project.py thành công",
                           ["Dòng log mong đợi", "Ý nghĩa"],
                           expected_outputs, col_widths=[8, 8])
    
    add_body_text(doc, "Thời gian chạy ước tính (máy tính phổ thông Intel i5, RAM 8GB):")
    time_estimates = [
        "Load Stops.csv (101MB): ~15–30 giây",
        "Merge và Transform: ~5–10 giây",
        "KMeans clustering: ~2–3 giây",
        "Tạo bản đồ HTML: ~10–15 giây",
        "Tổng cộng: ~30–60 giây",
    ]
    for idx, t in enumerate(time_estimates):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(3)
        letter = chr(97 + idx)
        run = p.add_run(f"({letter}) {t}")
        set_run_font(run, size=13)
    
    # Lưu file
    output_path = Path("c:/Users/My PC/Downloads/DE/PhanTich_Code_TfL_Nhom14.docx")
    doc.save(str(output_path))
    print(f"✅ Đã tạo FILE 2: {output_path}")
    return output_path


# ============================================================
# MAIN
# ============================================================
# HELPER FUNCTIONS FOR APPENDICES
# ============================================================

def add_appendix_c_content(doc):
    """Đọc tệp kiem_tra_va_kich_ban_thuyet_trinh.md và parse động phần 1 (Kiểm toán kỹ thuật) vào Word"""
    md_path = "kiem_tra_va_kich_ban_thuyet_trinh.md"
    if not os.path.exists(md_path):
        add_body_text(doc, "Tài liệu kiểm toán kỹ thuật chưa được tạo.")
        return
        
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()
    list_counter = 0
        
    in_code_block = False
    code_text = []
    
    in_table = False
    table_headers = []
    table_rows = []
    
    for line in lines:
        stripped = line.strip()
        
        # Dừng lại khi bắt đầu Phần 2 (Kịch bản thuyết trình) để tránh chèn kịch bản thuyết trình vào báo cáo chính thức
        if "PHẦN 2:" in stripped or "PHÂN 2:" in stripped:
            break
            
        # Xử lý code block
        if stripped.startswith("```"):
            if in_code_block:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.0)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run("\n".join(code_text))
                set_run_font(run, "Courier New", size=10)
                code_text = []
                in_code_block = False
            else:
                in_code_block = True
            continue
            
        if in_code_block:
            code_text.append(line)
            continue
            
        # Xử lý bảng
        if stripped.startswith("|"):
            parts = [p.strip() for p in stripped.split("|")[1:-1]]
            if not in_table:
                in_table = True
                table_headers = parts
            else:
                if not all(c == '-' or c == ':' for c in parts[0]):
                    table_rows.append(parts)
            continue
        else:
            if in_table:
                add_table_with_caption(doc, "Bảng phụ lục thông tin kiểm toán", table_headers, table_rows)
                in_table = False
                table_headers = []
                table_rows = []
                
        if not stripped:
            continue
            
        # Parse Markdown headings
        if stripped.startswith("# "):
            add_paragraph(doc, stripped[2:].upper(), align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_before=12, space_after=6)
        elif stripped.startswith("## "):
            add_heading_1(doc, "", stripped[3:])
        elif stripped.startswith("### "):
            add_heading_2(doc, "", stripped[4:])
        elif stripped.startswith("> "):
            add_paragraph(doc, stripped[2:], align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=13, italic=True, space_after=6)
        elif stripped.startswith("- ") or stripped.startswith("• "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_after = Pt(3)
            letter = chr(97 + list_counter)
            list_counter += 1
            r_pref = p.add_run(f"({letter}) ")
            set_run_font(r_pref, size=13)
            text = stripped[2:]
            if "**" in text:
                parts = text.split("**")
                for idx, part in enumerate(parts):
                    run = p.add_run(part)
                    set_run_font(run, size=13, bold=(idx % 2 == 1))
            else:
                run = p.add_run(text)
                set_run_font(run, size=13)
        elif not stripped:
            # Dòng trống, giữ nguyên list_counter
            continue
        else:
            # Dòng khác (heading, table, etc.) reset list_counter
            list_counter = 0
            # Đoạn văn thường
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.0)
            p.paragraph_format.space_after = Pt(6)
            text = line
            if "**" in text:
                parts = text.split("**")
                for idx, part in enumerate(parts):
                    run = p.add_run(part)
                    set_run_font(run, size=13, bold=(idx % 2 == 1))
            else:
                run = p.add_run(text)
                set_run_font(run, size=13)


def add_appendix_d_from_markdown(doc, md_path):
    """Đọc tệp markdown và parse động các heading, table, list, code block vào Word"""
    if not os.path.exists(md_path):
        add_body_text(doc, "Tài liệu kỹ thuật chưa được tạo.")
        return
        
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()
    list_counter = 0
        
    in_code_block = False
    code_text = []
    
    in_table = False
    table_headers = []
    table_rows = []
    
    for line in lines:
        stripped = line.strip()
        
        # Xử lý code block
        if stripped.startswith("```"):
            if in_code_block:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.0)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run("\n".join(code_text))
                set_run_font(run, "Courier New", size=10)
                code_text = []
                in_code_block = False
            else:
                in_code_block = True
            continue
            
        if in_code_block:
            code_text.append(line)
            continue
            
        # Xử lý bảng
        if stripped.startswith("|"):
            parts = [p.strip() for p in stripped.split("|")[1:-1]]
            if not in_table:
                in_table = True
                table_headers = parts
            else:
                if not all(c == '-' or c == ':' for c in parts[0]):
                    table_rows.append(parts)
            continue
        else:
            if in_table:
                add_table_with_caption(doc, "Bảng phụ lục thông tin kỹ thuật", table_headers, table_rows)
                in_table = False
                table_headers = []
                table_rows = []
                
        if not stripped:
            continue
            
        # Parse Markdown headings
        if stripped.startswith("# "):
            add_paragraph(doc, stripped[2:].upper(), align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_before=12, space_after=6)
        elif stripped.startswith("## "):
            add_heading_1(doc, "", stripped[3:])
        elif stripped.startswith("### "):
            add_heading_2(doc, "", stripped[4:])
        elif stripped.startswith("> "):
            add_paragraph(doc, stripped[2:], align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=13, italic=True, space_after=6)
        elif stripped.startswith("- ") or stripped.startswith("• "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_after = Pt(3)
            letter = chr(97 + list_counter)
            list_counter += 1
            r_pref = p.add_run(f"({letter}) ")
            set_run_font(r_pref, size=13)
            text = stripped[2:]
            if "**" in text:
                parts = text.split("**")
                for idx, part in enumerate(parts):
                    run = p.add_run(part)
                    set_run_font(run, size=13, bold=(idx % 2 == 1))
            else:
                run = p.add_run(text)
                set_run_font(run, size=13)
        elif not stripped:
            # Dòng trống, giữ nguyên list_counter
            continue
        else:
            # Dòng khác (heading, table, etc.) reset list_counter
            list_counter = 0
            # Đoạn văn thường
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.0)
            p.paragraph_format.space_after = Pt(6)
            text = line
            if "**" in text:
                parts = text.split("**")
                for idx, part in enumerate(parts):
                    run = p.add_run(part)
                    set_run_font(run, size=13, bold=(idx % 2 == 1))
            else:
                run = p.add_run(text)
                set_run_font(run, size=13)



# ============================================================

if __name__ == "__main__":
    print("=" * 60)
    print("TẠO 2 FILE BÁO CÁO CHO DỰ ÁN TfL LONDON — NHÓM 14")
    print("=" * 60)
    
    print("\n📄 Đang tạo FILE 1: BaoCao_TfL_London_Nhom14_Final.docx...")
    f1 = create_main_report()
    
    print("\n📄 Đang tạo FILE 2: PhanTich_Code_TfL_Nhom14.docx...")
    f2 = create_code_analysis()
    
    print("\n" + "=" * 60)
    print("✅ HOÀN THÀNH! Hai file đã được tạo:")
    print(f"   1. {f1}")
    print(f"   2. {f2}")
    print("=" * 60)


